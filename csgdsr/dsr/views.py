from django.shortcuts import render, redirect,get_object_or_404
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from django import forms
from .models import CheckPost, CSR, BNSSMissingCase, OtherCases,Other_Agencies, MaritimeAct, Officer, MPS, CheckPost,Other_Agencies, AttackOnTNFishermen_Choices, ArrestOfTNFishermen_Choices, ArrestOfSLFishermen_Choices, SeizedItemCategory, CustomUser, SeizedItemCategory, PS, RescueAtBeach,RescueAtSea,Seizure,Forecast,AttackOnTNFishermen, ArrestOfTNFishermen, ArrestOfSLFishermen,OnRoadVehicleStatus,OnWaterVehicleStatus,VVCmeeting,BeatDetails,Proforma,BoatPatrol,VehicleCheckPost,Atvpatrol,VehicleCheckothers, Unit, Headquarters, Zone, Range

from django.db.models import Count, Sum

from django.utils.dateparse import parse_date

from .forms import CustomSignupForm, UpdateUserForm,OfficerForm, CheckPostForm, CSRForm, BNSSMissingCaseForm,othercasesForm, MaritimeActForm,Other_AgenciesForm,OfficerForm, MPSForm, CheckPostForm,Other_AgenciesForm, AttackOnTNFishermen_ChoicesForm,ArrestOfTNFishermen_ChoicesForm, ArrestOfSLFishermen_ChoicesForm, SeizedItemCategoryForm,PSForm, RescueAtBeachForm,RescueAtSeaForm,SeizureForm,ForecastForm,AttackOnTNFishermenForm,ArrestOfTNFishermenForm,ArrestOfSLFishermenForm,OnRoadVehicleStatusForm, OnWaterVehicleStatusForm,  VVCmeetingForm, BeatDetailsForm, ProformaForm, BoatPatrolForm, VehicleCheckPostForm, AtvpatrolForm, VehicleCheckothersForm, HeadquartersForm, ZoneForm, RangeForm, UnitForm 


from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.urls import reverse_lazy, reverse
from django.contrib.auth.models import User
from django.utils import timezone

from django.contrib.auth.views import PasswordChangeView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import JsonResponse
from django.db.models import Q

from django.template.loader import render_to_string
from docx.shared import Inches
#home page
def home_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        remember = request.POST.get('remember')

        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            if not remember:
                request.session.set_expiry(0)  # session expires on browser close
            if user.is_superuser:
                return redirect('admin_dashboard')  # URL name for admin dashboard
            else:
                return redirect('user_dashboard')

        else:
            messages.error(request, 'Invalid username or password')

    return render(request, 'dsr/home.html')

# admin views
@login_required
def admin_dashboard_view(request):
    # buildings 
    total_mps = MPS.objects.count()
    total_checkpost = CheckPost.objects.count()
    total_officers = Officer.objects.count()

    # cases 
    total_csr = CSR.objects.count()
    total_bnss_cases = BNSSMissingCase.objects.count()
    total_other_cases= OtherCases.objects.count()
    total_maritime_cases = MaritimeAct.objects.count()

    # rescue 
    total_rescue_beach = RescueAtBeach.objects.count()
    total_rescued_beachvictims = RescueAtBeach.objects.aggregate(
        total=Sum('number_of_victims')
    )['total'] or 0 
    total_rescue_sea = RescueAtSea.objects.count()
    total_rescued_seavictims = RescueAtSea.objects.aggregate(
        total=Sum('number_of_victims')
    )['total'] or 0
    total_rescued_seaboats = RescueAtSea.objects.aggregate(
        total=Sum('number_of_boats_rescued')
    )['total'] or 0

    #tnfishermen arrest
    # --- CARD 3: ArrestOfTNFishermen stats ---
    arrest_qs = ArrestOfTNFishermen.objects.all()
    arrest_stats = arrest_qs.aggregate(
        total_entries=Count('id'),
        total_arrested=Sum('number_of_TNFishermen_arrested'),
        total_boats_seized=Sum('no_of_boats_seized'),
        total_released=Sum('number_of_TNFishermen_released'),
        total_boats_released=Sum('no_of_boats_released')
    )
    # --- CARD 4: AttackOnTNFishermen stats for SriLankan Navy ---
    navy = Other_Agencies.objects.filter(agency_name__iexact="SriLankan Navy").first()
    navy_stats = {}
    if navy:
        navy_stats = AttackOnTNFishermen.objects.filter(attacked_by=navy).aggregate(
            total_entries=Count('id'),
            total_injured=Sum('number_of_TNFishermen_injured'),
            total_missing=Sum('number_of_TNFishermen_missing'),
            total_died=Sum('number_of_TNFishermen_died')
        )
    # --- CARD 5: AttackOnTNFishermen stats for SriLankan Fishermen ---
    fishermen = Other_Agencies.objects.filter(agency_name__iexact="SriLankan Fishermen").first()
    fishermen_stats = {}
    if fishermen:
        fishermen_stats = AttackOnTNFishermen.objects.filter(attacked_by=fishermen).aggregate(
            total_entries=Count('id'),
            total_injured=Sum('number_of_TNFishermen_injured'),
            total_missing=Sum('number_of_TNFishermen_missing'),
            total_died=Sum('number_of_TNFishermen_died')
        )




    # sl fishermen arrest 
    total_slfisheremen_arrested_incident = ArrestOfSLFishermen.objects.count()
    total_slfisheremen_arrested = ArrestOfSLFishermen.objects.aggregate(total_arrested=Sum('number_of_SLFishermen_arrested'))
    ['total_arrested'] or 0
    total_slfisheremen_released = ArrestOfSLFishermen.objects.aggregate(total_released=Sum('number_of_SLFishermen_released'))['total_released'] or 0
    total_slboats_seized = ArrestOfSLFishermen.objects.aggregate(total_seized=Sum('no_of_boats_seized'))['total_seized'] or 0
    total_slboats_released = ArrestOfSLFishermen.objects.aggregate(total_released=Sum('no_of_boats_released'))['total_released'] or 0

    # attack on tn fishermen
    total_tnfisheremen_attacked_incident = AttackOnTNFishermen.objects.count()
    total_tnfishermen_injured = AttackOnTNFishermen.objects.aggregate(total_injured=Sum('number_of_TNFishermen_injured'))['total_injured'] or 0
    total_tnfisheremen_killed = AttackOnTNFishermen.objects.aggregate(total_killed=Sum('number_of_TNFishermen_died'))['total_killed'] or 0
    total_tnfishermen_missing= AttackOnTNFishermen.objects.aggregate(total_injured=Sum('number_of_TNFishermen_missing'))['total_injured'] or 0








    twelve_ton_boats = OnWaterVehicleStatus.objects.filter(boat_type='12_TON_BOAT')
    total_12t_boats= OnWaterVehicleStatus.objects.filter(boat_type='12_TON_BOAT').count()
    working_12ton_count = twelve_ton_boats.filter(working_status='WORKING').count()
    not_working_12ton_count = twelve_ton_boats.filter(working_status='NOT_WORKING').count()
    condemned_12ton_count = twelve_ton_boats.filter(working_status='CONDEMNED').count()

    five_ton_boats = OnWaterVehicleStatus.objects.filter(boat_type='5_TON_BOAT')
    total_5t_boats = OnWaterVehicleStatus.objects.filter(boat_type='5_TON_BOAT').count()
    working_5ton_count = five_ton_boats.filter(working_status='WORKING').count()
    not_working_5ton_count = five_ton_boats.filter(working_status='NOT_WORKING').count()
    condemned_5ton_count = five_ton_boats.filter(working_status='CONDEMNED').count()

    jetskis = OnWaterVehicleStatus.objects.filter(boat_type='JET_SKI')
    total_jetskis = OnWaterVehicleStatus.objects.filter(boat_type='JET_SKI').count()
    working_jetskis= jetskis.filter(working_status='WORKING').count()
    not_working_jetskis = jetskis.filter(working_status='NOT_WORKING').count()
    condemned_jetskis = jetskis.filter(working_status='CONDEMNED').count()

    jetboats= OnWaterVehicleStatus.objects.filter(boat_type='JET_BOAT')
    total_jetboats = OnWaterVehicleStatus.objects.filter(boat_type='JET_BOAT').count()
    working_jetboats = jetboats.filter(working_status='WORKING').count()
    not_working_jetboats = jetboats.filter(working_status='NOT_WORKING').count()
    condemned_jetboats = jetboats.filter(working_status='CONDEMNED').count()



    twowheelers = OnRoadVehicleStatus.objects.filter(vehicle_type='TWO_WHEELER')
    total_twowheelers = OnRoadVehicleStatus.objects.filter(vehicle_type='TWO_WHEELER').count()
    working_twowheelers = twowheelers.filter(working_status='WORKING').count()
    not_working_twowheelers = twowheelers.filter(working_status='NOT_WORKING').count()
    condemned_twowheelers = twowheelers.filter(working_status='CONDEMNED').count()

    fourwheelers = OnRoadVehicleStatus.objects.filter(vehicle_type='FOUR_WHEELER')
    total_fourwheelers = OnRoadVehicleStatus.objects.filter(vehicle_type='FOUR_WHEELER').count()
    working_fourwheelers = fourwheelers.filter(working_status='WORKING').count()
    not_working_fourwheelers = fourwheelers.filter(working_status='NOT_WORKING').count()
    condemned_fourwheelers = fourwheelers.filter(working_status='CONDEMNED').count()

    atvs= OnRoadVehicleStatus.objects.filter(vehicle_type='ATV')
    total_atv=OnRoadVehicleStatus.objects.filter(vehicle_type='ATV').count()
    working_atv = atvs.filter(working_status='WORKING').count()
    not_working_atv = atvs.filter(working_status='NOT_WORKING').count()
    condemned_atv = atvs.filter(working_status='CONDEMNED').count()

    total_vvc_meetings = VVCmeeting.objects.count()
    total_villagers_attended = VVCmeeting.objects.aggregate(total=Sum('number_of_villagers'))['total'] or 0

    total_boat_patrols = BoatPatrol.objects.count()
    total_boats_checked= BoatPatrol.objects.aggregate(total=Sum('numberof_boats_checked'))['total'] or 0
    total_atv_patrols = Atvpatrol.objects.count()
    total_vehicle_checkposts = VehicleCheckPost.objects.count()
    total_vehicle_check_others = VehicleCheckothers.objects.count()
    total_vehicles_checked_others = VehicleCheckothers.objects.aggregate(
        total=Sum('number_of_vehicles_checked')
    )['total'] or 0
    total_vehicles_checked_checkpost = VehicleCheckPost.objects.aggregate(total=Sum('number_of_vehicles_checked'))['total'] or 0


    total_daybeats = BeatDetails.objects.aggregate(total=Sum('day_beat_count'))['total'] or 0
    total_nightbeats = BeatDetails.objects.aggregate(total=Sum('night_beat_count'))['total'] or 0



    context = {
        'total_mps': total_mps,
        'total_checkpost': total_checkpost,
        'total_csr': total_csr,
        'total_bnss_cases': total_bnss_cases,
        'total_other_cases': total_other_cases,
        'total_maritime_cases': total_maritime_cases,
        'total_rescue_beach': total_rescue_beach,
        'total_rescued_beachvictims': total_rescued_beachvictims,
        'total_rescue_sea': total_rescue_sea,
        'total_rescued_seavictims': total_rescued_seavictims,
        'total_rescued_seaboats': total_rescued_seaboats,
        'total_12t_boats': total_12t_boats,
        'working_12ton_count': working_12ton_count,
        'not_working_12ton_count': not_working_12ton_count,
        'condemned_12ton_count': condemned_12ton_count,
        'total_5t_boats': total_5t_boats,
        'working_5ton_count': working_5ton_count,
        'not_working_5ton_count': not_working_5ton_count,
        'condemned_5ton_count': condemned_5ton_count,
        'total_jetskis': total_jetskis,
        'working_jetskis': working_jetskis,
        'not_working_jetskis': not_working_jetskis,
        'condemned_jetskis': condemned_jetskis,
        'total_jetboats': total_jetboats,
        'working_jetboats': working_jetboats,
        'not_working_jetboats': not_working_jetboats,
        'condemned_jetboats': condemned_jetboats,
        'total_twowheelers': total_twowheelers,
        'working_twowheelers': working_twowheelers,
        'not_working_twowheelers': not_working_twowheelers,
        'condemned_twowheelers': condemned_twowheelers,
        'total_fourwheelers': total_fourwheelers,
        'working_fourwheelers': working_fourwheelers,
        'not_working_fourwheelers': not_working_fourwheelers,
        'condemned_fourwheelers': condemned_fourwheelers,
        'total_atv': total_atv,
        'working_atv': working_atv,
        'not_working_atv': not_working_atv,
        'condemned_atv': condemned_atv,
        'total_boat_patrols': total_boat_patrols,
        'total_boats_checked': total_boats_checked,
        'total_atv_patrols': total_atv_patrols,
        'total_vehicle_checkposts': total_vehicle_checkposts,
        'total_vehicle_check_others': total_vehicle_check_others,
        'total_vehicles_checked_others': total_vehicles_checked_others,
        'total_vehicles_checked_checkpost': total_vehicles_checked_checkpost,
        'total_vvc_meetings': total_vvc_meetings,
        'total_villagers_attended': total_villagers_attended,
        'total_daybeats': total_daybeats,
        'total_nightbeats': total_nightbeats,
        # Arrest Stats
        'total_arrest_entries': arrest_stats['total_entries'],
        'number_of_TNFishermen_arrested': arrest_stats['total_arrested'],
        'no_of_boats_seized': arrest_stats['total_boats_seized'],
        'number_of_TNFishermen_released': arrest_stats['total_released'],
        'no_of_boats_released': arrest_stats['total_boats_released'],

        'total_slfisheremen_arrested_incident': total_slfisheremen_arrested_incident,
        'total_slfisheremen_arrested':total_slfisheremen_arrested,
        'total_slboats_seized': total_slboats_seized,
        'total_slfisheremen_released': total_slfisheremen_released,
        'total_slboats_released': total_slboats_released,

        
        # Attack Stats
        'navy_attack_data': navy_stats,
        'fishermen_attack_data': fishermen_stats,

    }    
    return render(request, 'dsr/admin/admin_dashboard.html',context)

@login_required
def admin_users_view(request, user_id=None):
    users = CustomUser.objects.all().order_by('username')
    user_instance = get_object_or_404(CustomUser, id=user_id) if user_id else None

    if request.method == 'POST':
        form = CustomSignupForm(request.POST, instance=user_instance)
        if form.is_valid():
            form.save()
            if user_instance:
                messages.success(request, "User updated successfully.")
            else:
                messages.success(request, "User created successfully.")
            return redirect('admin_users_page')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = CustomSignupForm(instance=user_instance)

    context = {
        'form': form,
        'users': users,
        'edit_mode': user_instance is not None,
        'user_id': user_instance.id if user_instance else '',
    }
    return render(request, 'dsr/admin/add_edituser.html', context)

@login_required
def admin_officers_strength_view(request, officer_id=None):
    officers = Officer.objects.all().order_by('rank')
    if officer_id:
        officer_instance = get_object_or_404(Officer, id=officer_id)
    else:
        officer_instance = None

    if request.method == 'POST':
        if officer_instance:
            form = OfficerForm(request.POST, instance=officer_instance)
            if form.is_valid():
                form.save()
                messages.success(request, "Officer details updated successfully!")
                return redirect('admin_officers_strength_page')
            else:
                messages.error(request, "Please correct the errors below.")
        else:
            form = OfficerForm(request.POST)
            if form.is_valid():
                form.save()
                messages.success(request, "Officer added successfully!")
                return redirect('admin_officers_strength_page')
            else:
                messages.error(request, "Please correct the errors below.")
    else:
        form = OfficerForm(instance=officer_instance)

    context = {
        'form': form,
        'officers': officers,
        'edit_mode': officer_instance is not None,
        'officer_id': officer_instance.id if officer_instance else '',
    }
    return render(request, 'dsr/admin/officers_details.html', context)

@login_required
def admin_MPS_buildings_view(request, mps_id=None, checkpost_id=None,ps_id=None, headquarters_id=None, zone_id=None, range_id=None, unit_id=None):
    mps_list = MPS.objects.all().order_by('name')
    checkpost_list = CheckPost.objects.all().order_by('name')
    ps_list = PS.objects.all().order_by('name')
    headquarters_list = Headquarters.objects.all().order_by('name')
    zone_list = Zone.objects.all().order_by('name')
    range_list = Range.objects.all().order_by('name')
    unit_list = Unit.objects.all().order_by('name')


    mps_instance = get_object_or_404(MPS, id=mps_id) if mps_id else None
    checkpost_instance = get_object_or_404(CheckPost, id=checkpost_id) if checkpost_id else None
    ps_instance = get_object_or_404(PS, id=ps_id) if ps_id else None
    headquarters_instance = get_object_or_404(Headquarters, id=headquarters_id) if headquarters_id else None
    zone_instance = get_object_or_404(Zone, id=zone_id) if zone_id else None
    range_instance = get_object_or_404(Range, id=range_id) if range_id else None
    unit_instance = get_object_or_404(Unit, id=unit_id) if unit_id else None

    if request.method == 'POST':
        if 'delete_mps' in request.POST:
            MPS.objects.filter(id=request.POST.get('delete_mps')).delete()
            messages.success(request, "MPS deleted successfully.")
            return redirect('admin_MPS_buildings_page')

        if 'delete_checkpost' in request.POST:
            CheckPost.objects.filter(id=request.POST.get('delete_checkpost')).delete()
            messages.success(request, "CheckPost deleted successfully.")
            return redirect('admin_MPS_buildings_page')

        if 'delete_ps' in request.POST:
            PS.objects.filter(id=request.POST.get('delete_ps')).delete()
            messages.success(request, "PS deleted successfully.")
            return redirect('admin_MPS_buildings_page')

        if 'delete_headquarters' in request.POST:
            Headquarters.objects.filter(id=request.POST.get('delete_headquarters')).delete()
            messages.success(request, "Headquarters deleted successfully.")
            return redirect('admin_MPS_buildings_page')

        if 'delete_zone' in request.POST:
            Zone.objects.filter(id=request.POST.get('delete_zone')).delete()
            messages.success(request, "Zone deleted successfully.")
            return redirect('admin_MPS_buildings_page')

        if 'delete_range' in request.POST:
            Range.objects.filter(id=request.POST.get('delete_range')).delete()
            messages.success(request, "Range deleted successfully.")
            return redirect('admin_MPS_buildings_page')

        if 'delete_unit' in request.POST:
            Unit.objects.filter(id=request.POST.get('delete_unit')).delete()
            messages.success(request, "Unit deleted successfully.")
            return redirect('admin_MPS_buildings_page')


        # Handle MPS Form
        if 'mps_submit' in request.POST:
            mps_form = MPSForm(request.POST, instance=mps_instance)
            if mps_form.is_valid():
                mps_form.save()
                if mps_instance:
                    messages.success(request, "MPS updated successfully.")
                else:
                    messages.success(request, "MPS added successfully.")
                return redirect('admin_MPS_buildings_page')
            else:
                messages.error(request, "Please correct the MPS form errors.")
            mps_form = MPSForm()
        # Handle CheckPost Form
        elif 'checkpost_submit' in request.POST:
            checkpost_form = CheckPostForm(request.POST, instance=checkpost_instance)
            if checkpost_form.is_valid():
                checkpost_form.save()
                if checkpost_instance:
                    messages.success(request, "Checkpost updated successfully.")
                else:
                    messages.success(request, "Checkpost added successfully.")
                return redirect('admin_MPS_buildings_page')
            else:
                messages.error(request, "Please correct the Checkpost form errors.")
            mps_form = MPSForm()  # Empty MPS form
        # Handle PS Form
        elif 'ps_submit' in request.POST:
            ps_form = PSForm(request.POST, instance=ps_instance)
            if ps_form.is_valid():
                ps_form.save()
                if ps_instance:
                    messages.success(request, "PS updated successfully.")
                else:
                    messages.success(request, "PS added successfully.")
                return redirect('admin_MPS_buildings_page')
            else:
                messages.error(request, "Please correct the PS form errors.")
            mps_form = MPSForm()
        
        elif 'headquarters_submit' in request.POST:
            headquarters_form = HeadquartersForm(request.POST, instance=headquarters_instance)
            if headquarters_form.is_valid():
                headquarters_form.save()
                if headquarters_instance:
                    messages.success(request, "Headquarters updated successfully.")
                else:
                    messages.success(request, "Headquarters added successfully.")
                return redirect('admin_MPS_buildings_page')
            else:
                messages.error(request, "Please correct the Headquarters form errors.")
            mps_form = MPSForm()
        
        elif 'zone_submit' in request.POST:
            zone_form = ZoneForm(request.POST, instance=zone_instance)
            if zone_form.is_valid():
                zone_form.save()
                if zone_instance:
                    messages.success(request, "Zone updated successfully.")
                else:
                    messages.success(request, "Zone added successfully.")
                return redirect('admin_MPS_buildings_page')
            else:
                messages.error(request, "Please correct the Zone form errors.")
            mps_form = MPSForm()

        elif 'range_submit' in request.POST:
            range_form = RangeForm(request.POST, instance=range_instance)
            if range_form.is_valid():
                range_form.save()
                if range_instance:
                    messages.success(request, "Range updated successfully.")
                else:
                    messages.success(request, "Range added successfully.")
                return redirect('admin_MPS_buildings_page')
            else:
                messages.error(request, "Please correct the Range form errors.")
            mps_form = MPSForm()
        
        elif 'unit_submit' in request.POST:
            unit_form = UnitForm(request.POST, instance=unit_instance)
            if unit_form.is_valid():
                unit_form.save()
                if unit_instance:
                    messages.success(request, "Unit updated successfully.")
                else:
                    messages.success(request, "Unit added successfully.")
                return redirect('admin_MPS_buildings_page')
            else:
                messages.error(request, "Please correct the Unit form errors.")
            mps_form = MPSForm()




    else:
        mps_form = MPSForm(instance=mps_instance)
        checkpost_form = CheckPostForm(instance=checkpost_instance)
        ps_form = PSForm(instance=ps_instance)
        headquarters_form = HeadquartersForm(instance=headquarters_instance)
        zone_form = ZoneForm(instance=zone_instance)
        range_form = RangeForm(instance=range_instance)
        unit_form = UnitForm(instance=unit_instance)
    context = {
        'mps_form': mps_form,
        'checkpost_form': checkpost_form,
        'ps_form': ps_form,
        'headquarters_form': headquarters_form,
        'zone_form': zone_form,
        'range_form': range_form,
        'unit_form': unit_form,
        'mps_list': mps_list,
        'checkpost_list': checkpost_list,
        'ps_list': ps_list,
        'headquarters_list': headquarters_list,
        'zone_list': zone_list,
        'range_list': range_list,
        'unit_list': unit_list,
        'edit_mps': mps_instance is not None,
        'edit_checkpost': checkpost_instance is not None,
        'edit_ps': ps_instance is not None,
        'edit_headquarters': headquarters_instance is not None,
        'edit_zone': zone_instance is not None,
        'edit_range': range_instance is not None,
        'edit_unit': unit_instance is not None,
        'mps_id': mps_instance.id if mps_instance else '',
        'checkpost_id': checkpost_instance.id if checkpost_instance else '',
        'ps_id': ps_instance.id if ps_instance else '',
        'headquarters_id': headquarters_instance.id if headquarters_instance else '',
        'zone_id': zone_instance.id if zone_instance else '',
        'range_id': range_instance.id if range_instance else '',
        'unit_id': unit_instance.id if unit_instance else '',
    }
    return render(request, 'dsr/admin/buildings.html', context)

@login_required
def admin_other_agencies_view(request, agency_id=None, attacker_id=None, arrest_tn_id=None, arrest_sl_id=None):
    
    agencies = Other_Agencies.objects.all().order_by('agency_name')
    # attackers = AttackOnTNFishermen_Choices.objects.all().order_by('attacker_name')
    # arrest_tn = ArrestOfTNFishermen_Choices.objects.all().order_by('arrested_by')
    # arrest_sl = ArrestOfSLFishermen_Choices.objects.all().order_by('arrested_by')

    agency_instance = get_object_or_404(Other_Agencies, id=agency_id) if agency_id else None
    # attacker_instance = get_object_or_404(AttackOnTNFishermen_Choices, id=attacker_id) if attacker_id else None
    # arrest_tn_instance = get_object_or_404(ArrestOfTNFishermen_Choices, id=arrest_tn_id) if arrest_tn_id else None
    # arrest_sl_instance = get_object_or_404(ArrestOfSLFishermen_Choices, id=arrest_sl_id) if arrest_sl_id else None

    if request.method == 'POST':
        
        # Other Agencies
        if 'agency_submit' in request.POST:
            form = Other_AgenciesForm(request.POST, instance=agency_instance)
            if form.is_valid():
                form.save()
                messages.success(request, "Other Agency saved successfully.")
                return redirect('admin_other_agencies_page')
            else:
                messages.error(request, "Please correct the errors in Other Agency form.")
            attacker_form = AttackOnTNFishermen_ChoicesForm()
            arrest_tn_form = ArrestOfTNFishermen_ChoicesForm()
            arrest_sl_form = ArrestOfSLFishermen_ChoicesForm()

        # Attack on TN Fishermen
        # elif 'attacker_submit' in request.POST:
        #     attacker_form = AttackOnTNFishermen_ChoicesForm(request.POST, instance=attacker_instance)
        #     if attacker_form.is_valid():
        #         attacker_form.save()
        #         messages.success(request, "Attacker name saved successfully.")
        #         return redirect('admin_other_agencies_page')
        #     else:
        #         messages.error(request, "Please correct the errors in Attacker form.")
        #     form = Other_AgenciesForm()
        #     arrest_tn_form = ArrestOfTNFishermen_ChoicesForm()
        #     arrest_sl_form = ArrestOfSLFishermen_ChoicesForm()

        # # Arrest of TN Fishermen
        # elif 'arrest_tn_submit' in request.POST:
        #     arrest_tn_form = ArrestOfTNFishermen_ChoicesForm(request.POST, instance=arrest_tn_instance)
        #     if arrest_tn_form.is_valid():
        #         arrest_tn_form.save()
        #         messages.success(request, "Arresting Authority (TN) saved successfully.")
        #         return redirect('admin_other_agencies_page')
        #     else:
        #         messages.error(request, "Please correct the errors in Arresting Authority (TN) form.")
        #     form = Other_AgenciesForm()
        #     attacker_form = AttackOnTNFishermen_ChoicesForm()
        #     arrest_sl_form = ArrestOfSLFishermen_ChoicesForm()

        # # Arrest of SL Fishermen
        # elif 'arrest_sl_submit' in request.POST:
        #     arrest_sl_form = ArrestOfSLFishermen_ChoicesForm(request.POST, instance=arrest_sl_instance)
        #     if arrest_sl_form.is_valid():
        #         arrest_sl_form.save()
        #         messages.success(request, "Arresting Authority (SL) saved successfully.")
        #         return redirect('admin_other_agencies_page')
        #     else:
        #         messages.error(request, "Please correct the errors in Arresting Authority (SL) form.")
        #     form = Other_AgenciesForm()
        #     attacker_form = AttackOnTNFishermen_ChoicesForm()
        #     arrest_tn_form = ArrestOfTNFishermen_ChoicesForm()

    else:
        form = Other_AgenciesForm(instance=agency_instance)
        # attacker_form = AttackOnTNFishermen_ChoicesForm(instance=attacker_instance)
        # arrest_tn_form = ArrestOfTNFishermen_ChoicesForm(instance=arrest_tn_instance)
        # arrest_sl_form = ArrestOfSLFishermen_ChoicesForm(instance=arrest_sl_instance)

    context = {
        'form': form,
        # 'attacker_form': attacker_form,
        # 'arrest_tn_form': arrest_tn_form,
        # 'arrest_sl_form': arrest_sl_form,
        'agencies': agencies,
        # 'attackers': attackers,
        # 'arrest_tn': arrest_tn,
        # 'arrest_sl': arrest_sl,
        'edit_agency': agency_instance is not None,
        # 'edit_attacker': attacker_instance is not None,
        # 'edit_arrest_tn': arrest_tn_instance is not None,
        # 'edit_arrest_sl': arrest_sl_instance is not None,
        'agency_id': agency_instance.id if agency_instance else '',
        # 'attacker_id': attacker_instance.id if attacker_instance else '',
        # 'arrest_tn_id': arrest_tn_instance.id if arrest_tn_instance else '',
        # 'arrest_sl_id': arrest_sl_instance.id if arrest_sl_instance else '',
    }
    return render(request, 'dsr/admin/other_agencies.html', context)

@login_required
def admin_seizure_items_view(request, item_id=None):
    items = SeizedItemCategory.objects.all().order_by('item_name')
    item_instance = get_object_or_404(SeizedItemCategory, id=item_id) if item_id else None

    if request.method == 'POST':
        form = SeizedItemCategoryForm(request.POST, instance=item_instance)
        if form.is_valid():
            form.save()
            if item_instance:
                messages.success(request, "Seized Item updated successfully.")
            else:
                messages.success(request, "Seized Item added successfully.")
            return redirect('admin_seizure_items_page')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = SeizedItemCategoryForm(instance=item_instance)

    context = {
        'form': form,
        'items': items,
        'edit_mode': item_instance is not None,
        'item_id': item_instance.id if item_instance else '',
    }
    return render(request, 'dsr/admin/seizure_items.html', context)

#add add checkppost view
@login_required
def add_checkpost(request):
    if request.method == 'POST':
        form = CheckPostForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "New checkpost added successfully!")
            return redirect('add_checkpost')  # redirect to the same form or elsewhere
    else:
        form = CheckPostForm()
    
    return render(request, 'dsr/admin.add_checkpost.html', {'form': form})

#Other_AgenciesForm view
@login_required
def other_agencies_form_view(request):
    if request.method == 'POST':
        form = Other_AgenciesForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Other agency added successfully!")
            return redirect('other_agencies_form')
    else:
        form = Other_AgenciesForm()
    return render(request, 'dsr/admin.other_agencies_form.html', {'form': form})

@login_required
def admin_vehicle_boat_view(request, vehicle_id=None, boat_id=None):
    vehicle_list = OnRoadVehicleStatus.objects.all().order_by('submitted_at')
    boat_list = OnWaterVehicleStatus.objects.all().order_by('submitted_at')
    

    vehicle_instance = get_object_or_404(OnRoadVehicleStatus, id=vehicle_id) if vehicle_id else None
    boat_instance = get_object_or_404(OnWaterVehicleStatus, id=boat_id) if boat_id else None
    
    if request.method == 'POST':
        
        # Handle vehicle Form
        if 'vehicle_submit' in request.POST:
            vehicle_form = OnRoadVehicleStatusForm(request.POST, instance=vehicle_instance)
            if vehicle_form.is_valid():
                vehicle_form.save()
                if vehicle_instance:
                    messages.success(request, "Vehicle Details updated successfully.")
                else:
                    messages.success(request, "Vehicle Details added successfully.")
                return redirect('admin_vehicle_boat_page')
            else:
                messages.error(request, "Please correct the Vehicle details form errors.")
            boat_form = OnWaterVehicleStatusForm()  # Empty Checkpost form

        # Handle CheckPost Form
        elif 'boat_submit' in request.POST:
            boat_form = OnWaterVehicleStatusForm(request.POST, instance=boat_instance)
            if boat_form.is_valid():
                boat_form.save()
                if boat_instance:
                    messages.success(request, "Boat Details updated successfully.")
                else:
                    messages.success(request, "Boat Details added successfully.")
                return redirect('admin_vehicle_boat_page')
            else:
                messages.error(request, "Please correct the Boat Details form errors.")
            
    else:
        vehicle_form = OnRoadVehicleStatusForm(instance=vehicle_instance)
        boat_form = OnWaterVehicleStatusForm(instance=boat_instance)
    context = {
        'vehicle_form': vehicle_form,
        'boat_form': boat_form,
        'vehicle_list': vehicle_list,
        'boat_list': boat_list,
        'edit_vehicle': vehicle_instance is not None,
        'edit_boat': boat_instance is not None,
        'vehicle_id': vehicle_instance.id if vehicle_instance else '',
        'boat_id': boat_instance.id if boat_instance else ''
    }
    return render(request, 'dsr/admin/admin_vehicle_boat_details.html', context)


def signup_view(request):
    if request.method == 'POST':
        form = CustomSignupForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Account created successfully. You can now log in.')
            return redirect('home')
    else:
        form = CustomSignupForm()
    return render(request, 'dsr/signup.html', {'form': form})

#admin reports
def admin_admin_cases_summary_view(request):
    csr_list = CSR.objects.all().order_by('-date_of_receipt')
    bnss_cases = BNSSMissingCase.objects.filter(case_category='194 BNSS').order_by('-date_of_receipt')
    missing_cases = BNSSMissingCase.objects.filter(case_category='Missing').order_by('-date_of_receipt')
    other_cases = OtherCases.objects.all().order_by('-date_of_receipt')
    maritime_cases= MaritimeAct.objects.all().order_by('-date_of_receipt')
    
    return render(request, 'dsr/admin/admin_cases_summary.html', {
        'csr_list': csr_list,
        'bnss_cases': bnss_cases,
        'missing_cases': missing_cases,
        'other_cases': other_cases,
        'maritime_cases': maritime_cases,

    })


def admin_vvc_beat_proforma_summary_view(request):
    vvc_records = VVCmeeting.objects.all().order_by('-submitted_at')
    beat_records = BeatDetails.objects.all().order_by('-submitted_at')
    proforma_records = Proforma.objects.all().order_by('-submitted_at')
    return render(request, 'dsr/admin/admin_proforma_summary.html', {
        'vvc_records': vvc_records,
        'beat_records': beat_records,
        'proforma_records':proforma_records
    })
    
    

def admin_patrol_check_summary_view(request):
    boatpatrol_records = BoatPatrol.objects.all().order_by('-submitted_at')
    atv_records = Atvpatrol.objects.all().order_by('-submitted_at')
    vehiclec_records = VehicleCheckPost.objects.all().order_by('-submitted_at')
    vehicleo_records = VehicleCheckothers.objects.all().order_by('-submitted_at')
    return render(request, 'dsr/admin/admin_vehicle_check_patrol_summary.html', {
        'boatpatrol_records': boatpatrol_records,
        'atv_records': atv_records,
        'vehiclec_records':vehiclec_records,
        'vehicleo_records':vehicleo_records
    })



def admin_fishermen_attack_arrest_summary_view(request):
    return render(request, 'dsr/admin/admin_fishermen_attack_arrest_summary.html')

def admin_assets_summary_view(request):
    return render(request, 'dsr/admin/admin_assets.html')


# users 
def user_dashboard_view(request):
    return render(request, 'dsr/user/user_dashboard.html')

def logout_view(request):
    logout(request)
    messages.success(request, "You have been logged out.")
    return redirect('home')  # Replace 'login' with the name or path to your login page

class CustomPasswordChangeView(LoginRequiredMixin, PasswordChangeView):
    template_name = 'dsr/registration/password_change.html'
    success_url = reverse_lazy('home')

    def form_valid(self, form):
        messages.success(self.request, "Your password is successfully updated.")
        return super().form_valid(form)

@login_required
def update_user_view(request): 
    if request.method == 'POST':
        form = UpdateUserForm(request.POST, instance=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, "Your account has been updated. Please log in again.")
            logout(request)  # Log the user out after updating
            return redirect('home')  # Replace 'home' with your desired URL name
    else:
        form = UpdateUserForm(instance=request.user)
    return render(request, 'dsr/registration/update_user.html', {'form': form})

@login_required
def forms_view(request):
    cards = [
        {"title": "CSR", "url_name": "csr_form", "icon": "bi-file-earmark-text", "color": "#0d6efd"},
        {"title": "194 BNSS", "url_name": "bnss_missing_form", "icon": "bi-clipboard-data", "color": "#6610f2"},
        {"title": "Missing", "url_name": "bnss_missing_form", "icon": "bi-person-dash", "color": "#dc3545"},
        {"title": "Maritime Act", "url_name": "maritimeact_form", "icon": "bi-globe", "color": "#20c997"},
        {"title": "Other Cases", "url_name": "othercases_form", "icon": "bi-briefcase", "color": "#6f42c1"},
        {"title": "Beach Rescue", "url_name": "rescue_at_beach_form", "icon": "bi-life-preserver", "color": "#fd7e14"},
        {"title": "Sea Rescue", "url_name": "rescue_at_sea_form", "icon": "bi-life-preserver", "color": "#fd7e14"},
        {"title": "Seizure", "url_name": "seizure_form", "icon": "bi-shield-check", "color": "#198754"},
        {"title": "Forecast", "url_name": "forecast_form", "icon": "bi-cloud-sun", "color": "#0dcaf0"},
        {"title": "Attack on TN Fishermen", "url_name": "attack_tnfishermen_form", "icon": "bi-person-x", "color": "#ffc107"},
        {"title": "TN Fishermen Arrest", "url_name": "arrest_tnfishermen_form", "icon": "bi-handcuffs", "color": "#dc3545"},
        {"title": "SL Fishermen Arrest", "url_name": "arrest_slfishermen_form", "icon": "bi-handcuffs", "color": "#dc3545"},
        {"title": "Vehicle Status", "url_name": "onroad_vehicle_status_form", "icon": "bi-truck", "color": "#6c757d"},
        {"title": "Boat Status", "url_name": "onwater_vehicle_status_form", "icon": "bi-truck", "color": "#6c757d"},
        {"title": "VVC", "url_name": "vvc_meeting_form", "icon": "bi-folder2-open", "color": "#20c997"},
        {"title": "Beat", "url_name": "beat_details_form", "icon": "bi-compass", "color": "#0d6efd"},
        {"title": "Proforma", "url_name": "proforma_form", "icon": "bi-ui-checks", "color": "#198754"},
        {"title": "Boat Patrol", "url_name": "boat_patrol_form", "icon": "bi-ship", "color": "#0dcaf0"},
        {"title": "ATV Patrol", "url_name": "atv_patrol_form", "icon": "bi-ship", "color": "#0dcaf0"},
        {"title": "Vehicle Check - CheckPost", "url_name": "vehicle_checkpost_form", "icon": "bi-search", "color": "#6f42c1"},
        {"title": "Vehicle Check - CheckPost", "url_name": "vehicle_check_others_form", "icon": "bi-search", "color": "#6f42c1"},
    ]
    return render(request, 'dsr/user/forms_page.html', {'cards': cards})

@login_required
def csr_form_view(request):
    if request.method == 'POST':
        form = CSRForm(request.POST)
        if form.is_valid():
            csr = form.save(commit=False)
            csr.user = request.user  # ✅ Assign the logged-in user
            csr.save()
            messages.success(request, 'CSR form submitted successfully!')
            return redirect('cases_summary')
    else:
        form = CSRForm()
    return render(request, 'dsr/user/forms/csr_form.html', {'form': form})

@login_required
def bnss_missing_form_view(request):
    if request.method == 'POST':
        form = BNSSMissingCaseForm(request.POST)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.user = request.user
            # instance.date_submitted = timezone.now()
            instance.save()
            messages.success(request, "Form submitted successfully!")
            return redirect('cases_summary')  # or wherever you want to go
        else:
            messages.error(request, "Please correct the errors below.")
    else:
        form = BNSSMissingCaseForm()

    return render(request, 'dsr/user/forms/194bnss_missing_form.html', {'form': form})

@login_required
def othercases_form(request):
    if request.method == 'POST':
        form = othercasesForm(request.POST)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.user = request.user
            instance.save()
            messages.success(request, "Form submitted successfully!")
            return redirect('cases_summary')  # or wherever you want to go
        else:
            messages.error(request, "Please correct the errors below.")
    else:
        form = othercasesForm()

    return render(request, 'dsr/user/forms/othercases_form.html',{'form': form})

@login_required
def maritimeact_form_view(request):
    if request.method == 'POST':
        form = MaritimeActForm(request.POST)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.user = request.user
            instance.save()
            messages.success(request, "Maritime Act form submitted successfully!")
            return redirect('cases_summary')
    else:
        form = MaritimeActForm()
    
    return render(request, 'dsr/user/forms/maritimeact_form.html', {'form': form})

@login_required
def rescue_at_beach_form_view(request, record_id=None):
    record = get_object_or_404(RescueAtBeach, id=record_id) if record_id else None

    if request.method == 'POST':
        form = RescueAtBeachForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            rescue_record = form.save(commit=False)
            rescue_record.user = request.user  # Track logged-in user
            rescue_record.save()
            if record:
                messages.success(request, "Rescue details updated successfully.")
            else:
                messages.success(request, "Rescue details submitted successfully.")
            return redirect('rescue_seizure_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = RescueAtBeachForm(instance=record)

    context = {
        'form': form,
        'edit_mode': record is not None,
    }
    return render(request, 'dsr/user/forms/rescuebeach_form.html', context)

@login_required
def rescue_at_sea_form_view(request,record_id=None):
    record = get_object_or_404(RescueAtSea, id=record_id) if record_id else None
    if request.method == 'POST':
        form = RescueAtSeaForm(request.POST)
        if form.is_valid():
            rescue_sea_record = form.save(commit=False)
            rescue_sea_record.user = request.user
            rescue_sea_record.save()
            if record:
                messages.success(request, "Rescue at Sea details updated successfully!")
            else:
                messages.success(request, "Rescue at Sea form submitted successfully!")
            return redirect('rescue_seizure_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = RescueAtSeaForm()
    return render(request, 'dsr/user/forms/rescueatsea_form.html', {'form': form})


@login_required
def seizure_form_view(request,record_id=None):
    record = get_object_or_404(Seizure, id=record_id) if record_id else None
    if request.method == 'POST':
        form = SeizureForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            seizure_record = form.save(commit=False)
            seizure_record.user = request.user
            seizure_record.save()
            if record:
                messages.success(request, "Seizure details updated successfully!")
            else:
                messages.success(request, "Seizure form submitted successfully!")
            return redirect('rescue_seizure_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = SeizureForm(instance=record)
    return render(request, 'dsr/user/forms/seizure_form.html', {'form': form})


@login_required
def forecast_form_view(request, record_id=None):
    record = get_object_or_404(Forecast, id=record_id) if record_id else None
    if request.method == 'POST':
        form = ForecastForm(request.POST)
        if form.is_valid():
            forecast_record = form.save(commit=False)
            forecast_record.user = request.user
            forecast_record.save()
            if record:
                messages.success(request, "Forecast details updated successfully!")
            else:
                messages.success(request, "Forecast form submitted successfully!")
            return redirect('forecast_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = ForecastForm(instance=record)
    return render(request, 'dsr/user/forms/forecast_form.html', {'form': form})

@login_required
def attack_tnfishermen_form(request, record_id=None):
    record= get_object_or_404(AttackOnTNFishermen, id=record_id) if record_id else None
    if request.method == 'POST':
        form = AttackOnTNFishermenForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            attack_record = form.save(commit=False)
            attack_record.user = request.user
            attack_record.save()
            if record:
                messages.success(request, "Attack on TN Fishermen details updated successfully!")
            else:
                messages.success(request, "Attack on TN Fishermen form submitted successfully!")
            return redirect('fishermen_attack_arrest_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = AttackOnTNFishermenForm(instance=record)
    context = {
        'form': form,
        'edit_mode': record is not None,
    }
    return render(request, 'dsr/user/forms/attack_tnfishermen_form.html', {'form': form})

@login_required
def arrest_tnfishermen_form(request, record_id=None):
    record = get_object_or_404(ArrestOfTNFishermen, id=record_id) if record_id else None
    if request.method == 'POST':
        form = ArrestOfTNFishermenForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            arrest_record = form.save(commit=False)
            arrest_record.user = request.user
            arrest_record.save()
            if record:
                messages.success(request, "Arrest of TN Fishermen details updated successfully!")
            else:
                messages.success(request, "Arrest of TN Fishermen form submitted successfully!")
            return redirect('fishermen_attack_arrest_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = ArrestOfTNFishermenForm(instance=record)
    context = {
        'form': form,
        'edit_mode': record is not None,
    }
    return render(request, 'dsr/user/forms/arrest_tnfishermen_form.html', context)
    
@login_required
def arrest_slfishermen_form(request, record_id=None):
    record = get_object_or_404(ArrestOfSLFishermen, id=record_id) if record_id else None
    if request.method == 'POST':
        form = ArrestOfSLFishermenForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            arrest_record = form.save(commit=False)
            arrest_record.user = request.user
            arrest_record.save()
            if record:
                messages.success(request, "Arrest of SL Fishermen details updated successfully!")
            else:
                messages.success(request, "Arrest of SL Fishermen form submitted successfully!")
            return redirect('fishermen_attack_arrest_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = ArrestOfSLFishermenForm(instance=record)
    context = {
        'form': form,
        'edit_mode': record is not None,
    }
    return render(request, 'dsr/user/forms/arrest_slfishermen_form.html', context)

@login_required
def onroad_vehicle_status_form_view(request, record_id=None):
    record = get_object_or_404(OnRoadVehicleStatus, id=record_id) if record_id else None
    if request.method == 'POST':
        form = OnRoadVehicleStatusForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            vehicle_record = form.save(commit=False)
            vehicle_record.user = request.user
            vehicle_record.save()
            if record:
                messages.success(request, "On Road Vehicle Status details updated successfully!")
            else:
                messages.success(request, "On Road Vehicle Status form submitted successfully!")
            return redirect('vehicle_status_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = OnRoadVehicleStatusForm(instance=record)
    context = {
        'form': form,
        'edit_mode': record is not None,
    }
    return render(request, 'dsr/user/forms/onroad_vehicle_form.html', context)

@login_required
def onwater_vehicle_status_form_view(request, record_id=None):
    record = get_object_or_404(OnWaterVehicleStatus, id=record_id) if record_id else None
    if request.method == 'POST':
        form = OnWaterVehicleStatusForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            vehicle_record = form.save(commit=False)
            vehicle_record.user = request.user
            vehicle_record.save()
            if record:
                messages.success(request, "On Water Vehicle Status details updated successfully!")
            else:
                messages.success(request, "On Water Vehicle Status form submitted successfully!")
            return redirect('vehicle_status_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = OnWaterVehicleStatusForm(instance=record)
    context = {
        'form': form,
        'edit_mode': record is not None,
    }
    return render(request, 'dsr/user/forms/onwater_vehicle_form.html', context)

@login_required
def vvc_meeting_form_view(request, record_id=None):
    record = get_object_or_404(VVCmeeting, id=record_id) if record_id else None
    if request.method == 'POST':
        form = VVCmeetingForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            vvc_record = form.save(commit=False)
            vvc_record.user = request.user
            vvc_record.save()
            if record:
                messages.success(request, "VVC Meeting details updated successfully!")
            else:
                messages.success(request, "VVC Meeting form submitted successfully!")
            return redirect('beat__vvc_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = VVCmeetingForm(instance=record)
    return render(request, 'dsr/user/forms/vvc_meeting_form.html', {'form': form})


@login_required
def beat_details_form_view(request, record_id=None):
    record = get_object_or_404(BeatDetails, id=record_id) if record_id else None
    if request.method == 'POST':
        form = BeatDetailsForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            beat_record = form.save(commit=False)
            beat_record.user = request.user
            beat_record.save()
            if record:
                messages.success(request, "Beat details updated successfully!")
            else:
                messages.success(request, "Beat details form submitted successfully!")
            return redirect('beat__vvc_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = BeatDetailsForm(instance=record)
    context = {
        'form': form,
        'edit_mode': record is not None,
    }
    return render(request, 'dsr/user/forms/beat_Details_form.html', context)

@login_required
def proforma_form_view(request, record_id=None):
    record = get_object_or_404(Proforma, id=record_id) if record_id else None
    if request.method == 'POST':
        form = ProformaForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            proforma_record = form.save(commit=False)
            proforma_record.user = request.user
            proforma_record.save()
            if record:
                messages.success(request, "Proforma details updated successfully!")
            else:
                messages.success(request, "Proforma form submitted successfully!")
            return redirect('proforma_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = ProformaForm(instance=record)
    context = {
        'form': form,
        'edit_mode': record is not None,
    }
    return render(request, 'dsr/user/forms/proforma_form.html', context)

@login_required
def boat_patrol_form_view(request, record_id=None):
    record = get_object_or_404(BoatPatrol, id=record_id) if record_id else None
    if request.method == 'POST':
        form = BoatPatrolForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            boat_patrol_record = form.save(commit=False)
            boat_patrol_record.user = request.user
            boat_patrol_record.save()
            if record:
                messages.success(request, "Boat Patrol details updated successfully!")
            else:
                messages.success(request, "Boat Patrol form submitted successfully!")
            return redirect('vehicle_check_patrol_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = BoatPatrolForm(instance=record)
    context = {
        'form': form,
        'edit_mode': record is not None,
    }
    return render(request, 'dsr/user/forms/boat_patrol_form.html', context)

@login_required
def atv_patrol_form_view(request, record_id=None):
    record = get_object_or_404(Atvpatrol, id=record_id) if record_id else None
    if request.method == 'POST':
        form = AtvpatrolForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            atv_patrol_record = form.save(commit=False)
            atv_patrol_record.user = request.user
            atv_patrol_record.save()
            if record:
                messages.success(request, "ATV Patrol details updated successfully!")
            else:
                messages.success(request, "ATV Patrol form submitted successfully!")
            return redirect('vehicle_check_patrol_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = AtvpatrolForm(instance=record)
    context = {
        'form': form,
        'edit_mode': record is not None,
    }
    return render(request, 'dsr/user/forms/atv_patrol_form.html', context)

@login_required
def vehicle_checkpost_form_view(request, record_id=None):
    record = get_object_or_404(VehicleCheckPost, id=record_id) if record_id else None
    if request.method == 'POST':
        form = VehicleCheckPostForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            vehicle_check_record = form.save(commit=False)
            vehicle_check_record.user = request.user
            vehicle_check_record.save()
            if record:
                messages.success(request, "Vehicle Check details updated successfully!")
            else:
                messages.success(request, "Vehicle Check form submitted successfully!")
            return redirect('vehicle_check_patrol_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = VehicleCheckPostForm(instance=record)
    context = {
        'form': form,
        'edit_mode': record is not None,
    }
    return render(request, 'dsr/user/forms/vehicle_checkport_form.html', context)

@login_required
def vehicle_check_others_form_view(request, record_id=None):
    record = get_object_or_404(VehicleCheckothers, id=record_id) if record_id else None
    if request.method == 'POST':
        form = VehicleCheckothersForm(request.POST, request.FILES, instance=record)
        if form.is_valid():
            vehicle_check_record = form.save(commit=False)
            vehicle_check_record.user = request.user
            vehicle_check_record.save()
            if record:
                messages.success(request, "Vehicle Check details updated successfully!")
            else:
                messages.success(request, "Vehicle Check form submitted successfully!")
            return redirect('vehicle_check_patrol_summary')
        else:
            messages.error(request, "Please correct the form errors.")
    else:
        form = VehicleCheckothersForm(instance=record)
    context = {
        'form': form,
        'edit_mode': record is not None,
    }
    return render(request, 'dsr/user/forms/vehicle_checkothers_form.html', context)

#submitted forms summary views
@login_required
def cases_registered_summary_view(request):
    csr_list = CSR.objects.filter(mps_limit__name=request.user.username).order_by('-date_of_receipt')

    bnss_cases = BNSSMissingCase.objects.filter(mps_limit__name=request.user.username, case_category='194 BNSS').order_by('-date_of_receipt')
    
    missing_cases = BNSSMissingCase.objects.filter(mps_limit__name=request.user.username, case_category='Missing').order_by('-date_of_receipt')
    
    other_cases = OtherCases.objects.filter(mps_limit__name=request.user.username).order_by('-date_of_receipt')
    
    maritime_cases= MaritimeAct.objects.filter(mps_limit__name=request.user.username).order_by('-date_of_receipt')
    
    return render(request, 'dsr/user/submitted_forms/cases_registered_summary.html', {
        'csr_list': csr_list,
        'bnss_cases': bnss_cases,
        'missing_cases': missing_cases,
        'other_cases': other_cases,
        'maritime_cases': maritime_cases,

    })

#rescue at beach summary view
@login_required
def rescue_beach_summary_view(request):
    rescue_records = RescueAtBeach.objects.order_by('-date_of_rescue')
    rescue_sea_record= RescueAtSea.objects.order_by('-date_of_rescue')
    seizure_records = Seizure.objects.order_by('-date_of_seizure')

    
    return render(request, 'dsr/user/submitted_forms/rescue_seizure_summary.html', {
        'rescue_records': rescue_records,
        'rescue_sea_record': rescue_sea_record,
        'seizure_records': seizure_records,
    })

# forecast summary view
@login_required
def forecast_summary_view(request):
    forecast_records = Forecast.objects.filter(mps_limit__name=request.user.username).order_by('-date_of_forecast')
    
    return render(request, 'dsr/user/submitted_forms/forecast_summary.html', {
        'forecast_records': forecast_records,
    })

#fishermen attack and arrest summary view
def fishermen_attack_arrest_summary(request):
    attacks = AttackOnTNFishermen.objects.filter(mps_limit__name=request.user.username).order_by('-submitted_at')
    tn_arrests = ArrestOfTNFishermen.objects.filter(mps_limit__name=request.user.username).order_by('-submitted_at')
    sl_arrests = ArrestOfSLFishermen.objects.filter(mps_limit__name=request.user.username).order_by('-submitted_at')
    return render(request, 'dsr/user/submitted_forms/fishermen_attack_arrest_summary.html', {
        'attacks': attacks,
        'tn_arrests': tn_arrests,
        'sl_arrests': sl_arrests,
    })

#vehicle status summary view
def vehicle_status_summary_view(request):
    onroad = OnRoadVehicleStatus.objects.all()
    onwater = OnWaterVehicleStatus.objects.all()

    return render(request, 'dsr/user/submitted_forms/vehicle_status_summary.html', {
        'onroad': onroad,
        'onwater': onwater,
    })

#vehicle status summary view
def beat__vvc_summary_view(request):
    vvc_records = VVCmeeting.objects.filter(mps_limit__name=request.user.username).order_by('-submitted_at')
    beat_records = BeatDetails.objects.filter(mps_limit__name=request.user.username).order_by('-submitted_at')

    return render(request, 'dsr/user/submitted_forms/beat_vvc_summary.html', {
        'vvc_records': vvc_records,
        'beat_records': beat_records,
    })

#proforma summary view
def proforma_summary_view(request):
    proforma_records = Proforma.objects.filter(user = request.user).order_by('-submitted_at')
    
    return render(request, 'dsr/user/submitted_forms/proforma_summary.html', {
        'proforma_records': proforma_records
    })

#patrol summary view
def vehicle_check_patrol_summary_view(request):
    boatpatrol_records = BoatPatrol.objects.filter(mps_limit__name=request.user.username).order_by('-submitted_at')
    atv_records = Atvpatrol.objects.filter(mps_limit__name=request.user.username).order_by('-submitted_at')
    vehiclec_records = VehicleCheckPost.objects.filter(mps_limit__name=request.user.username).order_by('-submitted_at')
    vehicleo_records = VehicleCheckothers.objects.filter(mps_limit__name=request.user.username).order_by('-submitted_at')
    
    return render(request, 'dsr/user/submitted_forms/vehicle_check_patrol_summary.html', {
        'boatpatrol_records': boatpatrol_records,'atv_records': atv_records,'vehiclec_records': vehiclec_records,'vehicleo_records': vehicleo_records
    })



#search for csr
@login_required
def csr_ajax_search_view(request):
    query = request.GET.get('q', '').strip()
    csr_list = CSR.objects.all()

    if query:
        csr_list = csr_list.filter(
            Q(csr_number__icontains=query) |
            Q(mps_limit__name__icontains=query) |
            Q(io__name__icontains=query) |
            Q(date_of_receipt__icontains=query)
            
            
              # Correct filter for Officer name
        )

    data = [
        {
            'id': csr.id,
            'csr_number': csr.csr_number,
            'mps_limit': str(csr.mps_limit),
            'date_of_receipt': csr.date_of_receipt.strftime('%d-%m-%Y'),
            
            'io': str(csr.io) if csr.io else '',  # Displays "Rank - Name"
        }
        for csr in csr_list
    ]
    
    return JsonResponse(data, safe=False)

#bnss search
@login_required
def bnss194_cases_ajax_search_view(request):
    query = request.GET.get('q', '').strip()
    cases = BNSSMissingCase.objects.filter(case_category='194 BNSS')

    if query:
        cases = cases.filter(
            Q(crime_number__icontains=query) |
            Q(date_of_occurrence__icontains=query) |
            Q(date_of_receipt__icontains=query) |
            Q(mps_limit__name__icontains=query) |
            Q(ps_limit__name__icontains=query) |  # Assuming PS model has station_name field
            Q(io__name__icontains=query) |
            Q(transfered_to__agency_name__icontains=query)
        )

    data = [
        {
            'id': case.id,
            'crime_number': case.crime_number,
            'mps_limit': str(case.mps_limit) if case.mps_limit else '',
            'ps_limit': str(case.ps_limit) if case.ps_limit else '',
            'date_of_occurrence': case.date_of_occurrence.strftime('%d-%m-%Y %H%Mhrs'),
            'date_of_receipt': case.date_of_receipt.strftime('%d-%m-%Y %H%Mhrs'),
            
            'io': str(case.io) if case.io else '',
            'transfered_to': str(case.transfered_to) if case.transfered_to else '',
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for missing cases
@login_required
def missing_ajax_search_view(request):
    query = request.GET.get('q', '').strip()
    cases = BNSSMissingCase.objects.filter(case_category='Missing')

    if query:
        cases = cases.filter(
            Q(crime_number__icontains=query) |
            Q(mps_limit__name__icontains=query) |
            Q(ps_limit__name__icontains=query) |

            Q(date_of_occurrence__icontains=query) |
            Q(date_of_receipt__icontains=query) |
            Q(io__name__icontains=query) |
            Q(transfered_to__agency_name__icontains=query)
        )

    data = [
        {
            'id': case.id,
            'crime_number': case.crime_number,
            'mps_limit': str(case.mps_limit) if case.mps_limit else '',
            'ps_limit': str(case.ps_limit) if case.ps_limit else '',
            'date_of_occurrence': case.date_of_occurrence.strftime('%d-%m-%Y %H%Mhrs'),
            'date_of_receipt': case.date_of_receipt.strftime('%d-%m-%Y %H%Mhrs'),
            'io': str(case.io) if case.io else '',
            'transfered_to': str(case.transfered_to) if case.transfered_to else '-',
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for other cases
@login_required
def othercases_ajax_search_view(request):
    query = request.GET.get('q', '').strip()
    cases = OtherCases.objects.all()

    if query:
        cases = cases.filter(
            Q(crime_number__icontains=query) |
            Q(mps_limit__name__icontains=query) |
            Q(ps_limit__name__icontains=query) |

            Q(date_of_occurrence__icontains=query) |
            Q(date_of_receipt__icontains=query) |
            Q(io__name__icontains=query) |
            Q(transfered_to__agency_name__icontains=query)
        )

    data = [
        {
            'id': case.id,
            'crime_number': case.crime_number,
            'mps_limit': str(case.mps_limit) if case.mps_limit else '',
            'ps_limit': str(case.ps_limit) if case.ps_limit else '',
            'date_of_occurrence': case.date_of_occurrence.strftime('%d-%m-%Y %H%Mhrs'),
            'date_of_receipt': case.date_of_receipt.strftime('%d-%m-%Y %H%Mhrs'),
            'ps_limit': case.ps_limit.name if case.ps_limit else '',
            'io': str(case.io) if case.io else '',
            'transfered_to': str(case.transfered_to) if case.transfered_to else '-',
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for maritime act cases
@login_required
def maritimeact_ajax_search_view(request):
    query = request.GET.get('q', '').strip()
    cases = MaritimeAct.objects.all()

    if query:
        cases = cases.filter(
            Q(crime_number__icontains=query) |
            Q(mps_limit__name__icontains=query) |
            Q(ps_limit__name__icontains=query) |
            Q(date_of_occurrence__icontains=query) |
            Q(date_of_receipt__icontains=query) |
            Q(io__name__icontains=query) |
            Q(transfered_to__agency_name__icontains=query)
            
        )

    data = [
        {
            'id': case.id,
            'crime_number': case.crime_number,
            'mps_limit': str(case.mps_limit) if case.mps_limit else '',
            'ps_limit': str(case.ps_limit) if case.ps_limit else '',
            'date_of_occurrence': case.date_of_occurrence.strftime('%d-%m-%Y %H%Mhrs'),
            'date_of_receipt': case.date_of_receipt.strftime('%d-%m-%Y %H%Mhrs'),
            'ps_limit': case.ps_limit.name if case.ps_limit else '',
            'io': str(case.io) if case.io else '',
            'transfered_to': str(case.transfered_to) if case.transfered_to else '-',
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for rescue at beach cases
@login_required
def rescue_at_beach_ajax_search_view(request):
    query = request.GET.get('q', '').strip()
    cases = RescueAtBeach.objects.filter(user=request.user)

    if query:
        cases = cases.filter(
            Q(date_of_rescue__icontains=query) |
            Q(place_of_rescue__icontains=query) |
            Q(number_of_victims__icontains=query)
        )

    data = [
        {
            'id': case.id,
            'date_of_rescue': case.date_of_rescue.strftime('%d-%m-%Y %H%Mhrs'),
            'place_of_rescue': case.place_of_rescue,
            'number_of_victims': case.number_of_victims,
            'rescue_beach_image': case.rescue_beach_image.url if case.rescue_beach_image else ''
   
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for rescue at sea cases
@login_required
def rescue_at_sea_ajax_search_view(request):    
    query = request.GET.get('q', '').strip()
    cases = RescueAtSea.objects.filter(user=request.user)

    if query:
        cases = cases.filter(
            Q(date_of_rescue__icontains=query) |
            Q(place_of_rescue__icontains=query) |
            Q(number_of_victims__icontains=query) 
        )

    data = [
        {
            'id': case.id,
            'date_of_rescue': case.date_of_rescue.strftime('%d-%m-%Y %H%Mhrs'),
            'place_of_rescue': case.place_of_rescue,
            'number_of_victims': case.number_of_victims,
            'rescue_sea_image': case.rescue_sea_image.url if case.rescue_sea_image else ''
   
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for seizure cases
@login_required
def seizure_ajax_search_view(request):
    query = request.GET.get('q', '').strip()
    cases = Seizure.objects.filter(user=request.user)

    if query:
        cases = cases.filter(
            Q(date_of_seizure__icontains=query) |
            Q(place_of_seizure__icontains=query) |
            Q(seized_item__item_name__icontains=query) |
            Q(quantity__icontains=query) |
            Q(handed_over_to__agency_name__icontains=query) 
        )

    data = [
        {
            'id': case.id,
            'date_of_seizure': case.date_of_seizure.strftime('%d-%m-%Y %H%Mhrs'),
            'place_of_seizure': case.place_of_seizure,
            'seized_item': str(case.seized_item) if case.seized_item else '',
            'quantity': case.quantity,
            'handed_over_to': str(case.handed_over_to) if case.handed_over_to else '',
            'seizure_image': case.seizure_image.url if case.seizure_image else ''
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for forecast cases
@login_required
def forecast_ajax_search_view(request):
    query = request.GET.get('q', '').strip()
    cases = Forecast.objects.filter(user=request.user)

    if query:
        cases = cases.filter(
            Q(date_of_forecast__icontains=query) |
            Q(place_of_forecast__icontains=query)
        )

    data = [
        {
            'id': case.id,
            'date_of_forecast': case.date_of_forecast.strftime('%d-%m-%Y'),
            'place_of_forecast': case.place_of_forecast,
            'forecast_details': case.forecast_details,
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for fishermen attack cases
def ajax_search_attack_tnfishermen(request):
    query = request.GET.get('q', '').strip()
    cases = AttackOnTNFishermen.objects.filter(user=request.user)

    if query:
        cases = cases.filter(
            Q(date_of_attack__icontains=query) |
            Q(place_of_attack__icontains=query) |
            Q(attacked_by__agency_name__icontains=query) |
            Q(district__icontains=query) 
        )

    data = [
        {
            'id': case.id,
            'date_of_attack': case.date_of_attack.strftime('%d-%m-%Y %H%Mhrs'),
            'place_of_attack': case.place_of_attack,
            'attacked_by': str(case.attacked_by) if case.attacked_by else '',
            
            'injured': str(case.number_of_TNFishermen_injured) if case.number_of_TNFishermen_injured else '',
            'missing': str(case.number_of_TNFishermen_missing) if case.number_of_TNFishermen_missing else '',
            'died': str(case.number_of_TNFishermen_died) if case.number_of_TNFishermen_died else '',
            'district': case.district,
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for fishermen arrest cases
def ajax_search_arrest_tnfishermen(request):
    query = request.GET.get('q', '').strip()
    cases = ArrestOfTNFishermen.objects.filter(user=request.user)

    if query:
        cases = cases.filter(
            Q(date_of_arrest__icontains=query) |
            Q(place_of_arrest__icontains=query) |
            Q(arrested_by__agency_name__icontains=query) |
            Q(district__icontains=query) 
        )

    data = [
        {
            'id': case.id,
            'date_of_arrest': case.date_of_arrest.strftime('%d-%m-%Y %H%Mhrs'),
            'place_of_arrest': case.place_of_arrest,
            'arrested_by': str(case.arrested_by) if case.arrested_by else '',
            'number_of_TNFishermen_arrested': str(case.number_of_TNFishermen_arrested) if case.number_of_TNFishermen_arrested else '',
            'no_of_boats_seized': str(case.no_of_boats_seized) if case.no_of_boats_seized else '',
            'district': case.district,
            'number_of_TNFishermen_released': str(case.number_of_TNFishermen_released) if case.number_of_TNFishermen_released else '',
            'no_of_boats_released': str(case.no_of_boats_released) if case.no_of_boats_released else '',
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for arrest of SL fishermen cases
def ajax_search_arrest_slfishermen(request):
    query = request.GET.get('q', '').strip()
    # cases = ArrestOfSLFishermen.objects.filter(user=request.user)
    user = request.user
    if user.is_superuser:
        cases = ArrestOfSLFishermen.objects.all()
    else:
        cases = ArrestOfSLFishermen.objects.filter(mps_limit__name=request.user.username)

    if query:
        cases = cases.filter(
            Q(date_of_arrest__icontains=query) |
            Q(place_of_arrest__icontains=query) |
            Q(arrested_by__agency_name__icontains=query) 
        )

    data = [
        {
            'id': case.id,
            'date_of_arrest': case.date_of_arrest.strftime('%d-%m-%Y %H%Mhrs'),
            'place_of_arrest': case.place_of_arrest,
            'arrested_by': str(case.arrested_by) if case.arrested_by else '',
            'number_of_SLFishermen_arrested': str(case.number_of_SLFishermen_arrested) if case.number_of_SLFishermen_arrested else '',
            'no_of_boats_seized': str(case.no_of_boats_seized) if case.no_of_boats_seized else '',
            'number_of_SLFishermen_released': str(case.number_of_SLFishermen_released) if case.number_of_SLFishermen_released else '',
            'no_of_boats_released': str(case.no_of_boats_released) if case.no_of_boats_released else '',
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for vvc details 
@login_required
def vvc_ajax_search_view(request):    
    query = request.GET.get('q', '').strip()
    # cases = VVCmeeting.objects.all()
    user = request.user
    if user.is_superuser:
        cases = VVCmeeting.objects.all()
    else:
        cases = VVCmeeting.objects.filter(mps_limit__name=request.user.username)

    if query:
        cases = cases.filter(
            Q(date_of_vvc__icontains=query) |
            Q(mps_limit__name__icontains=query) |
            Q(village_name__icontains=query) 
            
            
        )

    data = [
        {
            'id': case.id,
            'date_of_vvc': case.date_of_vvc.strftime('%d-%m-%Y'),
            'mps_limit': str(case.mps_limit) if case.mps_limit else '',
            'village_name':case.village_name,
            'number_of_villagers':case.number_of_villagers,
            'vvc_image':case.vvc_image.url if case.vvc_image else ''
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)




#search for beat details 
@login_required
def beat_ajax_search_view(request):    
    query = request.GET.get('q', '').strip()
    # cases = BeatDetails.objects.all()
    user = request.user
    if user.is_superuser:
        cases = BeatDetails.objects.all()
    else:
        cases = BeatDetails.objects.filter(mps_limit__name=request.user.username)
    

    if query:
        cases = cases.filter(
            Q(date_of_beat__icontains=query) |
            Q(mps_limit__name__icontains=query) 
        )

    data = [
        {
            'id': case.id,
            'date_of_beat': case.date_of_beat.strftime('%d-%m-%Y'),
            'mps_limit': str(case.mps_limit) if case.mps_limit else '',
            'day_beat_count': case.day_beat_count,
            'night_beat_count': case.night_beat_count
   
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for proforma 
@login_required
def proforma_ajax_search_view(request):    
    query = request.GET.get('q', '').strip()
    cases = Proforma.objects.all()

    if query:
        cases = cases.filter(
            Q(date_of_proforma__icontains=query) |
            Q(officer__name__icontains=query)

        )

    data = [
        {
            'id': case.id,
            'officer': str(case.officer) if case.officer else '',
            'date_of_proforma': case.date_of_proforma.strftime('%d-%m-%Y'),
            'mps_visited': case.mps_visited,
            'check_post_checked': case.check_post_checked,
            'boat_guard_checked': case.boat_guard_checked,
            'vvc_meeting_conducted': case.vvc_meeting_conducted,
            'villages_visited': case.villages_visited,
            'meetings_attended': case.meetings_attended,
            'awareness_programs_conducted': case.awareness_programs_conducted,
            'coastal_security_exercises_conducted': case.coastal_security_exercises_conducted,
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for onroad vehicle status

@login_required
def onroad_vehicle_status_ajax_search_view(request):    
    query = request.GET.get('q', '').strip()
    cases = OnRoadVehicleStatus.objects.all()

    if query:
        cases = cases.filter(
            Q(vehicle_type__icontains=query) |
            Q(vehicle_number__icontains=query) |
            Q(working_status__icontains=query) 
        )

    data = [
        {
            'id': case.id,
            'vehicle_type': case.get_vehicle_type_display(),
            'vehicle_number': case.vehicle_number,
            'alloted_to':case.alloted_to,
            'working_status': case.working_status
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for onwater vehicle status
@login_required
def onwater_vehicle_status_ajax_search_view(request):    
    query = request.GET.get('q', '').strip()
    cases = OnWaterVehicleStatus.objects.all()

    if query:
        cases = cases.filter(
            Q(boat_type__icontains=query) |
            Q(boat_number__icontains=query) |
            Q(working_status__icontains=query) 
        )

    data = [
        {
            'id': case.id,
            'boat_type': case.get_boat_type_display(),
            'boat_number': case.boat_number,
            'working_status': case.working_status
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for boat patrol

@login_required
def boat_patrol_ajax_search(request):    
    query = request.GET.get('q', '').strip()
    
    # cases = BoatPatrol.objects.all()  # <-- FIXED
    user = request.user
    if user.is_superuser:
        cases = BoatPatrol.objects.all()
    else:
        cases = BoatPatrol.objects.filter(mps_limit__name=request.user.username)



    if query:
        cases = cases.filter(
            Q(date_of_patrol__icontains=query) |
            Q(mps_limit__name__icontains=query) |
            Q(boat_type__icontains=query) |
            Q(boat_number__boat_number__icontains=query) |
            Q(numberof_boats_checked__icontains=query)
        )

    data = [
        {
            'id': case.id,
            'date_of_patrol': case.date_of_patrol.strftime('%d-%m-%Y'),  # format fixed
            'mps_limit': str(case.mps_limit) if case.mps_limit else '',
            'patrol_place': case.patrol_place,
            'boat_type': case.get_boat_type_display(),
            'boat_number': case.boat_number.boat_number,
            'numberof_boats_checked': case.numberof_boats_checked
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for atv patrol
@login_required
def atv_patrol_ajax_search_view(request):    
    query = request.GET.get('q', '').strip()
    # cases = Atvpatrol.objects.all()
    user = request.user
    if user.is_superuser:
        cases = Atvpatrol.objects.all()
    else:
        cases = Atvpatrol.objects.filter(mps_limit__name=request.user.username)

    if query:
        cases = cases.filter(
            Q(date_of_patrol__icontains=query) |
            Q(mps_limit__name__icontains=query) |
            Q(patrol_place__icontains=query) |
            Q(atv_number__vehicle_number__icontains=query) |
            Q(atv_number__vehicle_type__icontains=query)
        )

    data = [
        {
            'id': case.id,
            'date_of_patrol': case.date_of_patrol.strftime('%d-%m-%Y'),
            'mps_limit': str(case.mps_limit) if case.mps_limit else '',
            'patrol_place': case.patrol_place,
            'atv_number': (
                f"{case.atv_number.get_vehicle_type_display()} - {case.atv_number.vehicle_number}"
                if case.atv_number else "—"
            )
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)


#search for Vehicle check- checkpost
@login_required
def vehicle_checkpost_ajax_search_view(request):    
    query = request.GET.get('q', '').strip()
    # cases = VehicleCheckPost.objects.all()
    user = request.user
    if user.is_superuser:
        cases = VehicleCheckPost.objects.all()
    else:
        cases = VehicleCheckPost.objects.filter(mps_limit__name=request.user.username)


    if query:
        cases = cases.filter(
            Q(date_of_check__icontains=query) |
            Q(mps_limit__name__icontains=query) |
            Q(check_post__name__icontains=query) 
            
        )

    data = [
        {
            'id': case.id,
            'date_of_check': case.date_of_check.strftime('%d-%m-%Y'),
            'mps_limit': str(case.mps_limit) if case.mps_limit else '',
            'vehicle_check_start_time': case.vehicle_check_start_time.strftime('%H:%M') if case.vehicle_check_start_time else '',
            'vehicle_check_end_time': case.vehicle_check_end_time.strftime('%H:%M') if case.vehicle_check_end_time else '',
            'check_post': case.check_post.name ,
            'number_of_vehicles_checked': case.number_of_vehicles_checked,
            
            
   
        }
        for case in cases
    ]
    return JsonResponse(data, safe=False)

#search for Vehicle check- others
@login_required
def vehicle_check_others_ajax_search_view(request):    
    query = request.GET.get('q', '').strip()
    # cases = VehicleCheckothers.objects.all()
    user = request.user
    if user.is_superuser:
        cases = VehicleCheckothers.objects.all()
    else:
        cases = VehicleCheckothers.objects.filter(mps_limit__name=request.user.username)

    

    if query:

        cases = cases.filter(
            Q(date_of_check__icontains=query) |
            Q(mps_limit__name__icontains=query) |
            Q(place_of_check__icontains=query) 
        )

    data = [
    {
        'id': case.id,
        'date_of_check': case.date_of_check.strftime('%d-%m-%Y'),
        'mps_limit': str(case.mps_limit) if case.mps_limit else '',
        'vehicle_check_start_time': case.vehicle_check_start_time.strftime('%H:%M') if case.vehicle_check_start_time else '',
        'vehicle_check_end_time': case.vehicle_check_end_time.strftime('%H:%M') if case.vehicle_check_end_time else '',
        'place_of_check': case.place_of_check,
        'number_of_vehicles_checked': case.number_of_vehicles_checked,
    }
    for case in cases
    ]
    return JsonResponse(data, safe=False)

#exprt CSR to Word document
@login_required
def csr_export_word_view(request):
    csrs = CSR.objects.all().order_by('date_of_receipt')
    doc = Document()

    for csr in csrs:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        # Fields to include in the Word document
        fields = [
            ("CSR No.", csr.csr_number),
            ("MPS Limit ", csr.mps_limit),  # Directly use it as string
            ("Date of Receipt", csr.date_of_receipt.strftime('%Y-%m-%d') if csr.date_of_receipt else ''),
            ("Place of Occurrence", csr.place_of_occurrence),
            ("Petitioner", csr.petitioner),
            ("Counter Petitioner", csr.counter_petitioner),
            ("Nature of Petition", csr.nature_of_petition),
            ("Gist of Petition", csr.gist_of_petition),
            ("IO", csr.io)
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()  # Empty line
        doc.add_paragraph()  # Second empty line

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="CSR_Export.docx"'
    doc.save(response)
    return response

#export 194 BNSS cases to Word document
@login_required
def bnss_194_export_word_view(request):
    bnsss = BNSSMissingCase.objects.filter(case_category='194 BNSS').order_by('date_of_receipt')
    doc = Document()

    for bnss in bnsss:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        # Fields to include in the Word document
        fields = [
            ("Crime No.", bnss.crime_number),
            ("Police Station Limit", bnss.ps_limit),  # Directly use it as string
            ("MPS Limit", bnss.mps_limit),
            ("Date of Occurrence", bnss.date_of_occurrence.strftime('%Y-%m-%d %H:%M') if bnss.date_of_occurrence else ''),
            ("Date of Receipt", bnss.date_of_receipt.strftime('%Y-%m-%d') if bnss.date_of_receipt else ''),
            ("Place of Occurrence", bnss.place_of_occurrence),
            ("Petitioner", bnss.petitioner),
            ("Diseased", bnss.diseased if bnss.case_category == '194 BNSS' else ''),
            ("Missing Person", bnss.missing_person if bnss.case_category == 'Missing' else ''),
            ("Gist of Case", bnss.gist_of_case),
            ("IO", bnss.io),
            ("Transfered To", bnss.transfered_to),
            
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()  # Empty line
        doc.add_paragraph()  # Second empty line

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="BNSS194_Export.docx"'
    doc.save(response)
    return response

#export missing cases to Word document
@login_required
def missing_export_word_view(request):
    missings = BNSSMissingCase.objects.filter(case_category='Missing').order_by('date_of_receipt')
    doc = Document()

    for missing in missings:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        # Fields to include in the Word document
        fields = [
            ("Crime No.", missing.crime_number),
            ("Police Station Limit", missing.ps_limit),  # Directly use it as string
            ("MPS Limit", missing.mps_limit),
            ("Date of Occurrence", missing.date_of_occurrence.strftime('%Y-%m-%d %H:%M') if missing.date_of_occurrence else ''),
            ("Date of Receipt", missing.date_of_receipt.strftime('%Y-%m-%d') if missing.date_of_receipt else ''),
            ("Place of Occurrence", missing.place_of_occurrence),
            ("Petitioner", missing.petitioner),
            ("Diseased", missing.diseased if missing.case_category == '194 BNSS' else ''),
            ("Missing Person", missing.missing_person if missing.case_category == 'Missing' else ''),
            ("Gist of Case", missing.gist_of_case),
            ("IO", missing.io),
            ("Transfered To", missing.transfered_to),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()  # Empty line
        doc.add_paragraph()  # Second empty line

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Missing_Export.docx"'
    doc.save(response)
    return response

#export other cases to Word document
@login_required
def othercases_export_word_view(request):
    other_cases = OtherCases.objects.filter(user=request.user).order_by('date_of_receipt')
    doc = Document()

    for case in other_cases:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        # Fields to include in the Word document
        fields = [
            ("Crime No.", case.crime_number),
            ("Police Station Limit", case.ps_limit),  # Directly use it as string
            ("MPS Limit", case.mps_limit),
            ("Date of Occurrence", case.date_of_occurrence.strftime('%Y-%m-%d %H:%M') if case.date_of_occurrence else ''),
            ("Date of Receipt", case.date_of_receipt.strftime('%Y-%m-%d') if case.date_of_receipt else ''),
            ("Place of Occurrence", case.place_of_occurrence),
            ("Petitioner", case.petitioner),
            ("Diseased", case.diseased if case.diseased else ''),
            ("Injured", case.injured if case.injured else ''),
            ("Accused", case.accused if case.accused else ''),
            ("Gist of Case", case.gist_of_case),
            ("IO", case.io),
            ("Transfered To", case.transfered_to),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()  # Second empty line
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Other_Cases_Export.docx"'
    doc.save(response)
    return response

#export maritime act cases to Word document
@login_required
def maritimeact_export_word_view(request):
    maritime_cases = MaritimeAct.objects.filter(user=request.user).order_by('date_of_receipt')
    doc = Document()

    for case in maritime_cases:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        # Fields to include in the Word document
        fields = [
            ("Crime No.", case.crime_number),
            ("Police Station Limit", case.ps_limit),  # Directly use it as string
            ("MPS Limit", case.mps_limit),
            ("Date of Occurrence", case.date_of_occurrence.strftime('%Y-%m-%d %H:%M') if case.date_of_occurrence else ''),
            ("Date of Receipt", case.date_of_receipt.strftime('%Y-%m-%d') if case.date_of_receipt else ''),
            ("Place of Occurrence", case.place_of_occurrence),
            ("Petitioner", case.petitioner),
            ("Accused", case.accused if case.accused else ''),
            ("Gist of Case", case.gist_of_case),
            ("IO", case.io),
            ("Transfered To", case.transfered_to)
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()  # Second empty line
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Maritime_Cases_Export.docx"'
    doc.save(response)
    return response

#export rescue at beach cases to Word document
@login_required
def rescue_at_beach_export_word_view(request):
    rescueb_events = RescueAtBeach.objects.filter(user=request.user).order_by('date_of_rescue')
    doc = Document()

    for case in rescueb_events:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        fields = [
            ("Police Station Limit", case.ps_limit),
            ("MPS Limit", case.mps_limit),
            ("Date of Rescue", case.date_of_rescue.strftime('%Y-%m-%d %H:%M') if case.date_of_rescue else ''),
            ("Place of Rescue", case.place_of_rescue),
            ("Number of Persons Rescued", case.number_of_victims),
            ("Victim Details", case.victim_name),
            ("Rescuer Names", case.rescuer_name if case.rescuer_name else ''),
            ("Gist of Rescue", case.gist_of_rescue),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()

        # Add Rescue Image if available
        if case.rescue_beach_image:
            doc.add_paragraph("Rescue Image:")

            try:
                image_path = case.rescue_beach_image.path
                doc.add_picture(image_path, width=Inches(2.0))  # Adjust width as needed
            except Exception as e:
                doc.add_paragraph(f"Image could not be loaded: {e}")

        doc.add_paragraph()
        doc.add_paragraph()  # Extra space between entries

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Rescue_at_Beach.docx"'
    doc.save(response)
    return response

#export rescue at sea cases to Word document
@login_required
def rescue_at_sea_export_word_view(request):
    rescue_sea_events = RescueAtSea.objects.filter(user=request.user).order_by('date_of_rescue')
    doc = Document()

    for case in rescue_sea_events:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [

            ("MPS Limit", case.mps_limit),
            ("Date of Rescue", case.date_of_rescue.strftime('%Y-%m-%d %H:%M') if case.date_of_rescue else ''),
            ("Place of Rescue", case.place_of_rescue),
            ("Number of Persons Rescued", case.number_of_victims),
            ("Victim Details", case.victim_name),
            ("Number of Boats REscued", case.number_of_boats_rescued),
            ("Rescuer Names", case.rescuer_name if case.rescuer_name else ''),
            ("Gist of Rescue", case.gist_of_rescue),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()

        # Add Rescue Image if available
        if case.rescue_sea_image:
            doc.add_paragraph("Rescue Image:")

            try:
                image_path = case.rescue_sea_image.path
                doc.add_picture(image_path, width=Inches(2.0))  # Adjust width as needed
            except Exception as e:
                doc.add_paragraph(f"Image could not be loaded: {e}")

        doc.add_paragraph()
        doc.add_paragraph()
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Rescue_at_Sea.docx"'
    doc.save(response)
    return response

#export seizure cases to Word document
@login_required
def seizure_export_word_view(request):
    seizure_events = Seizure.objects.filter(user=request.user).order_by('date_of_seizure')
    doc = Document()

    for case in seizure_events:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Date of Seizure", case.date_of_seizure.strftime('%Y-%m-%d %H:%M') if case.date_of_seizure else ''),
            ("Place of Seizure", case.place_of_seizure),
            ("Seized Item", str(case.seized_item) if case.seized_item else ''),
            ("Quantity", case.quantity),
            ("Lattitude", case.latitude if case.latitude else ''),
            ("Longitude", case.longitude if case.longitude else ''),
            ("PS Limit", case.ps_limit.name if case.ps_limit else ''),
            ("MPS Limit", case.mps_limit.name if case.mps_limit else ''),
            ('Accused', case.accused if case.accused else ''),
            ("Handed Over To", str(case.handed_over_to) if case.handed_over_to else ''),
            ("Gist of Seizure", case.gist_of_seizure),
            ("Seizure Image", case.seizure_image.url if case.seizure_image else '')
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()

        # Add Seizure Image if available
        if case.seizure_image:
            doc.add_paragraph("Seizure Image:")

            try:
                image_path = case.seizure_image.path
                doc.add_picture(image_path, width=Inches(2.0))  # Adjust width as needed
            except Exception as e:
                doc.add_paragraph(f"Image could not be loaded: {e}")

        doc.add_paragraph()
        doc.add_paragraph()
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Seizure_Export.docx"'
    doc.save(response)
    return response

#export forecast cases to Word document
def forecast_export_word_view(request):
    forecast_events = Forecast.objects.filter(user=request.user).order_by('date_of_forecast')
    doc = Document()

    for case in forecast_events:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Date of Forecast", case.date_of_forecast.strftime('%Y-%m-%d %H:%M') if case.date_of_forecast else ''),
            ("Place of Forecast", case.place_of_forecast),
            ("Forecast Details", case.forecast_details),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Forecast_Export.docx"'
    doc.save(response)
    return response

#export tn fishermen attack cases to Word document
@login_required
def attack_tnfishermen_export_word_view(request):
    attack_cases = AttackOnTNFishermen.objects.filter(user=request.user).order_by('date_of_attack')
    doc = Document()

    for case in attack_cases:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("MPS Limit", case.mps_limit),
            ("Date of Attack", case.date_of_attack.strftime('%Y-%m-%d %H:%M') if case.date_of_attack else ''),
            ("Place of Attack", case.place_of_attack),
            ("Attacked By", str(case.attacked_by) if case.attacked_by else ''),
            ("District", case.district),
            ("Number of TN Fishermen Injured", str(case.number_of_TNFishermen_injured) if case.number_of_TNFishermen_injured else ''),
            ("Number of TN Fishermen Missing", str(case.number_of_TNFishermen_missing) if case.number_of_TNFishermen_missing else ''),
            ("Number of TN Fishermen Died", str(case.number_of_TNFishermen_died) if case.number_of_TNFishermen_died else ''),
            ("Victims", case.victim_names if case.victim_names else ''),
            ("Attacker Details", case.attacker_details if case.attacker_details else ''),
            ("Gist of Attack", case.gist_of_attack),
            ("trasnfered to", str(case.transfered_to) if case.transfered_to else ''),  
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="TN_Fishermen_Attack_Export.docx"'
    doc.save(response)
    return response

#export tn fishermen arrest cases to Word document
@login_required
def arrest_tnfishermen_export_word_view(request):
    arrest_cases = ArrestOfTNFishermen.objects.filter(user=request.user).order_by('date_of_arrest')
    doc = Document()

    for case in arrest_cases:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Date of Arrest", case.date_of_arrest.strftime('%Y-%m-%d %H:%M') if case.date_of_arrest else ''),
            ("Place of Arrest", case.place_of_arrest),
            ("Arrested By", str(case.arrested_by) if case.arrested_by else ''),
            ("District", case.district),
            ("MPS Limit", case.mps_limit),
            ("Number of TN Fishermen Arrested", str(case.number_of_TNFishermen_arrested) if case.number_of_TNFishermen_arrested else ''),
            ("fishermen Details", case.arrested_Fishermen_names if case.arrested_Fishermen_names else ''),
            ("Number of Boat Seized", str(case.no_of_boats_seized) if case.no_of_boats_seized else ''),
            ("Gist of Arrest", case.gist_of_arrest),
            ("Number of TN Fishermen Released", str(case.number_of_TNFishermen_released) if case.number_of_TNFishermen_released else ''),
            ("Number of Baots Released", case.no_of_boats_released if case.no_of_boats_released else '')            
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="TN_Fishermen_Arrest_Export.docx"'
    doc.save(response)
    return response

#export SL fishermen arrest cases to Word document
@login_required
def arrest_slfishermen_export_word_view(request):
    arrest_cases = ArrestOfSLFishermen.objects.filter(user=request.user).order_by('date_of_arrest')
    doc = Document()

    for case in arrest_cases:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Date of Arrest", case.date_of_arrest.strftime('%Y-%m-%d %H:%M') if case.date_of_arrest else ''),
            ("Place of Arrest", case.place_of_arrest),
            ("MPS Limit", case.mps_limit),
            ("Arrested By", str(case.arrested_by) if case.arrested_by else ''),
            ("Number of SL Fishermen Arrested", str(case.number_of_SLFishermen_arrested) if case.number_of_SLFishermen_arrested else ''),
            ("SL Fishermen Details", case.arrested_Fishermen_name if case.arrested_Fishermen_name else ''),
            ("Number of Boats Seized", str(case.no_of_boats_seized) if case.no_of_boats_seized else ''),
            ("Gist of Arrest", case.gist_of_arrest),
            
            ("Gist of Arrest", case.gist_of_arrest),
            ("Number of SL Fishermen Released", str(case.number_of_SLFishermen_released) if case.number_of_SLFishermen_released else ''),
            ("Number of Boats Released", case.no_of_boats_released if case.no_of_boats_released else ''),
            
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="SL_Fishermen_Arrest_Export.docx"'
    doc.save(response)
    return response


#export onroad vehicle status
@login_required
def onroad_vehicle_status_export_word_view(request):
    vehicles = OnRoadVehicleStatus.objects.filter(mps_limit__name=request.user.username).order_by('vehicle_type')
    doc = Document()

    for vehicle in vehicles:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Vehicle Type", vehicle.get_vehicle_type_display()),
            ("Vehicle Number", vehicle.vehicle_number),
            ("Alloted to",vehicle.alloted_to),
            ("Working Status", vehicle.get_working_status_display()),
            ("MPS Limit", str(vehicle.mps_limit) if vehicle.mps_limit else ''),
            ("Remarks", vehicle.remarks if vehicle.remarks else ''),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="OnRoad_Vehicle_Status.docx"'
    doc.save(response)
    return response

#export boat status as word
@login_required
def onwater_vehicle_status_export_word_view(request):
    boats = OnWaterVehicleStatus.objects.filter(mps_limit__name=request.user.username).order_by('boat_type')
    doc = Document()

    for boat in boats:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Boat Type", boat.get_boat_type_display()),
            ("Boat Number", boat.boat_number if boat.boat_number else ''),
            ("Working Status", boat.get_working_status_display()),
            ("MPS Limit", str(boat.mps_limit) if boat.mps_limit else ''),
            ("Remarks", boat.remarks if boat.remarks else ''),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="OnWater_Vehicle_Status.docx"'
    doc.save(response)
    return response

#vvc meeting export word
@login_required
def vvc_export_word_view(request):
    meetings = VVCmeeting.objects.filter(mps_limit__name=request.user.username).order_by('-date_of_vvc')
    doc = Document()

    for meeting in meetings:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Date of VVC Meeting", meeting.date_of_vvc.strftime('%d-%m-%Y') if meeting.date_of_vvc else ''),
            ("Village Name", meeting.village_name),
            ("Number of Villagers Participated", str(meeting.number_of_villagers) if meeting.number_of_villagers else ''),
            ("Conducted By", meeting.conducted_by),
            ("MPS Limit", str(meeting.mps_limit) if meeting.mps_limit else ''),
            ("VVC Details", meeting.vvc_details if meeting.vvc_details else ''),
            ("VVC Image", meeting.vvc_image.url if meeting.vvc_image else '')
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        # Add Seizure Image if available
        if meeting.vvc_image:
            doc.add_paragraph("VVC Image:")

            try:
                image_path = meeting.vvc_image.path
                doc.add_picture(image_path, width=Inches(2.0))  # Adjust width as needed
            except Exception as e:
                doc.add_paragraph(f"Image could not be loaded: {e}")

        doc.add_paragraph()
        doc.add_paragraph()
        

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="VVC_Meeting_Export.docx"'
    doc.save(response)
    return response

# beat exort word view
@login_required
def beat_export_word_view(request):
    beats = BeatDetails.objects.filter(mps_limit__name=request.user.username).order_by('-date_of_beat')
    doc = Document()

    for beat in beats:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Date of Beat", beat.date_of_beat.strftime('%d-%m-%Y') if beat.date_of_beat else ''),
            ("MPS Limit", str(beat.mps_limit) if beat.mps_limit else ''),
            ("Day Beat Count", str(beat.day_beat_count)),
            ("Night Beat Count", str(beat.night_beat_count)),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Beat_Details_Export.docx"'
    doc.save(response)
    return response

#proforma export word 
@login_required
def proforma_export_word_view(request):
    proformas = Proforma.objects.filter(user=request.user).order_by('-date_of_proforma')
    doc = Document()

    for pro in proformas:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Date of Proforma", pro.date_of_proforma.strftime('%d-%m-%Y') if pro.date_of_proforma else ''),
            ("MPS Visited", str(pro.mps_visited)),
            ("Check Posts Checked", str(pro.check_post_checked)),
            ("Boat Guards Checked", str(pro.boat_guard_checked)),
            ("VVC Meetings Conducted", str(pro.vvc_meeting_conducted)),
            ("Villages Visited", str(pro.villages_visited)),
            ("Meetings Attended", str(pro.meetings_attended)),
            ("Awareness Programs Conducted", str(pro.awareness_programs_conducted)),
            ("Coastal Security Exercises Conducted", str(pro.coastal_security_exercises_conducted)),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Proforma_Export.docx"'
    doc.save(response)
    return response

#boat patrol export view
@login_required
def boat_patrol_export_word_view(request):
    patrols = BoatPatrol.objects.filter(mps_limit__name=request.user.username).order_by('-date_of_patrol')
    doc = Document()

    for patrol in patrols:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Date of Patrol", patrol.date_of_patrol.strftime('%d-%m-%Y') if patrol.date_of_patrol else ''),
            ("Patrol Officer", str(patrol.patrol_officer) if patrol.patrol_officer else ''),
            ("Boat Type", patrol.get_boat_type_display()),
            ("Boat Number", patrol.boat_number if patrol.boat_number else ''),
            ("Patrol Time", f"{patrol.patrol_start_time.strftime('%H:%M')} - {patrol.patrol_end_time.strftime('%H:%M')} hrs"),
            ("Place of Patrol", patrol.patrol_place),
            ("MPS Limit", str(patrol.mps_limit) if patrol.mps_limit else ''),
            ("Number of Boats Checked", str(patrol.numberof_boats_checked)),
            ("Registration Number(s) of Boats Checked", patrol.registration_numberofboats_checked if patrol.registration_numberofboats_checked else ''),
            ("Remarks", patrol.remarks if patrol.remarks else ''),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Boat_Patrol_Export.docx"'
    doc.save(response)
    return response

#atv patrol export word
@login_required
def atv_patrol_export_word_view(request):
    patrols = Atvpatrol.objects.filter(mps_limit__name=request.user.username).order_by('-date_of_patrol')
    doc = Document()

    for patrol in patrols:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Date of Patrol", patrol.date_of_patrol.strftime('%d-%m-%Y') if patrol.date_of_patrol else ''),
            ("Patrol Officer", patrol.patrol_officer),
            ("ATV Number", patrol.atv_number if patrol.atv_number else ''),
            ("Patrol Time", f"{patrol.patrol_start_time.strftime('%H:%M')} - {patrol.patrol_end_time.strftime('%H:%M')} hrs"),
            ("Place of Patrol", patrol.patrol_place),
            ("MPS Limit", str(patrol.mps_limit) if patrol.mps_limit else ''),
            ("Remarks", patrol.remarks if patrol.remarks else ''),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="ATV_Patrol_Export.docx"'
    doc.save(response)
    return response

#vehicle check -post export word 
@login_required
def vehicle_checkpost_export_word_view(request):
    checks = VehicleCheckPost.objects.filter(mps_limit__name=request.user.username).order_by('-date_of_check')
    doc = Document()

    for check in checks:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Date of Check", check.date_of_check.strftime('%d-%m-%Y') if check.date_of_check else ''),
            ("Officer", str(check.officer) if check.officer else ''),
            ("Time of Check", f"{check.vehicle_check_start_time.strftime('%H:%M')} - {check.vehicle_check_end_time.strftime('%H:%M')} hrs"),
            ("Check Post", str(check.check_post)),
            ("MPS Limit", str(check.mps_limit) if check.mps_limit else ''),
            ("Number of Vehicles Checked", str(check.number_of_vehicles_checked)),
            ("Registration Numbers", check.registration_numbers if check.registration_numbers else ''),
            ("Remarks", check.remarks if check.remarks else ''),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Vehicle_Check_Post_Export.docx"'
    doc.save(response)
    return response

#vehicle check others export word view
@login_required
def vehicle_check_others_export_word_view(request):
    checks = VehicleCheckothers.objects.filter(mps_limit__name=request.user.username).order_by('-date_of_check')
    doc = Document()

    for check in checks:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        fields = [
            ("Date of Check", check.date_of_check.strftime('%d-%m-%Y') if check.date_of_check else ''),
            ("Officer", str(check.officer) if check.officer else ''),
            ("Time of Check", f"{check.vehicle_check_start_time.strftime('%H:%M')} - {check.vehicle_check_end_time.strftime('%H:%M')} hrs"),
            ("Place of Check", check.place_of_check),
            ("MPS Limit", str(check.mps_limit) if check.mps_limit else ''),
            ("Number of Vehicles Checked", str(check.number_of_vehicles_checked)),
            ("Registration Numbers", check.registration_numbers if check.registration_numbers else ''),
            ("Remarks", check.remarks if check.remarks else ''),
        ]

        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="Vehicle_Check_Others_Export.docx"'
    doc.save(response)
    return response


#export individual CSR
@login_required
def csr_download_view(request, pk):
    csr = get_object_or_404(CSR, pk=pk)

    # Create a new Word document
    doc = Document()
    doc.add_heading('Community Service Register (CSR)', level=1)

    # Create a table with 2 columns
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    # Helper function to add a row
    def add_row(field, value):
        row = table.add_row().cells
        row[0].text = field
        row[1].text = str(value) if value else ''

    # Add CSR details to the table
    add_row('CSR Number', csr.csr_number)
    add_row('MPS Limit ', str(csr.mps_limit))
    add_row('Date of Receipt', csr.date_of_receipt.strftime('%Y-%m-%d'))
    add_row('Place of Occurrence', csr.place_of_occurrence)
    add_row('Petitioner', csr.petitioner)
    add_row('Counter Petitioner', csr.counter_petitioner)
    add_row('Nature of Petition', csr.nature_of_petition)
    add_row('Gist of Petition', csr.gist_of_petition)
    add_row('IO', csr.io)

    # Prepare HTTP response with docx content
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=CSR_{csr.csr_number}.docx'

    doc.save(response)
    return response

#export individual BNSS194 case
@login_required
def bnss194_download_view(request, pk):
    case = get_object_or_404(BNSSMissingCase, pk=pk)

    doc = Document()
    doc.add_heading(f"{case.case_category} Case Details", level=1)

    # Create table with two columns
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    # Add data
    add_row('Crime Number', case.crime_number)
    add_row('Case Category', case.case_category)
    add_row('Police Station Limit', case.ps_limit)
    add_row('MPS Limit', case.mps_limit)
    add_row('Date of Occurrence', case.date_of_occurrence.strftime('%d-%m-%Y %H%Mhrs') if case.date_of_occurrence else '')
    add_row('Date of Receipt', case.date_of_receipt.strftime('%d-%m-%Y %H%Mhrs') if case.date_of_receipt else '')
    add_row('Place of Occurrence', case.place_of_occurrence)
    if case.case_category == '194 BNSS':
        add_row('Diseased', case.diseased)
    elif case.case_category == 'Missing':
        add_row('Missing Person', case.missing_person)
    add_row('Petitioner', case.petitioner)
    add_row('Gist of Case', case.gist_of_case)
    add_row('IO', case.io)
    add_row('Transfered To', case.transfered_to)

    # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"{case.case_category.replace(' ', '_')}_{case.crime_number}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

#export individual missing case
@login_required
def othercases_download_view(request, pk):
    case = get_object_or_404(OtherCases, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('Other Cases Details', level=1)

    # Create table with two columns
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    # Add data
    add_row('Crime Number', case.crime_number)
    add_row('Police Station Limit', case.ps_limit)
    add_row('MPS Limit', case.mps_limit)
    add_row('Date of Occurrence', case.date_of_occurrence.strftime('%d-%m-%Y %H%Mhrs') if case.date_of_occurrence else '')
    add_row('Date of Receipt', case.date_of_receipt.strftime('%d-%m-%Y %H%Mhrs') if case.date_of_receipt else '')
    add_row('Place of Occurrence', case.place_of_occurrence)
    add_row('Petitioner', case.petitioner)
    add_row('Diseased', case.diseased if case.diseased else '')
    add_row('Injured', case.injured if case.injured else '')
    add_row('Accused', case.accused if case.accused else '')
    add_row('Gist of Case', case.gist_of_case)
    add_row('IO', case.io)
    add_row('Transfered To', case.transfered_to)

    # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Other_Case_{case.crime_number}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

#export individual maritime act case
@login_required
def maritimeact_download_view(request, pk):
    case = get_object_or_404(MaritimeAct, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('Maritime Act Case Details', level=1)

    # Create table with two columns
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    # Add data
    add_row('Crime Number', case.crime_number)
    add_row('Police Station Limit', case.ps_limit)
    add_row('MPS Limit', case.mps_limit)
    add_row('Date of Occurrence', case.date_of_occurrence.strftime('%d-%m-%Y %H%Mhrs') if case.date_of_occurrence else '')
    add_row('Date of Receipt', case.date_of_receipt.strftime('%d-%m-%Y %H%Mhrs') if case.date_of_receipt else '')
    add_row('Place of Occurrence', case.place_of_occurrence)
    add_row('Petitioner', case.petitioner)
    add_row('Accused', case.accused if case.accused else '')
    add_row('Gist of Case', case.gist_of_case)
    add_row('IO', case.io)
    add_row('Transfered To', case.transfered_to)

    # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Maritime_Case_{case.crime_number}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

#export individual rescue at beach case
@login_required
def rescue_at_beach_download_view(request, pk):
    case = get_object_or_404(RescueAtBeach, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('Rescue at Beach Details', level=1)

    # Create table with two columns
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    # Add data
    add_row('Police Station Limit', case.ps_limit)
    add_row('MPS Limit', case.mps_limit)
    add_row('Date of Rescue', case.date_of_rescue.strftime('%d-%m-%Y %H:%M') if case.date_of_rescue else '')
    add_row('Place of Rescue', case.place_of_rescue)
    add_row('Number of Persons Rescued', case.number_of_victims)
    add_row('Victim Details', case.victim_name)
    add_row('Rescuer Names', case.rescuer_name if case.rescuer_name else '')
    add_row('Gist of Rescue', case.gist_of_rescue)

    doc.add_paragraph()

    # Add Rescue Image if available
    if case.rescue_beach_image:
        doc.add_paragraph("Rescue Image:")

        try:
            image_path = case.rescue_beach_image.path
            doc.add_picture(image_path, width=Inches(2.0))  # Adjust width as needed
        except Exception as e:
            doc.add_paragraph(f"Image could not be loaded: {e}")

    # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Rescue_at_Beach_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

#export individual rescue at sea case
@login_required
def rescue_at_sea_download_view(request, pk):
    case = get_object_or_404(RescueAtSea, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('Rescue at Sea Details', level=1)

    # Create table with two columns
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    # Add data
    add_row('MPS Limit', case.mps_limit)
    add_row('Date of Rescue', case.date_of_rescue.strftime('%d-%m-%Y %H:%M') if case.date_of_rescue else '')
    add_row('Place of Rescue', case.place_of_rescue)
    add_row('Number of Persons Rescued', case.number_of_victims)
    add_row('Victim Details', case.victim_name)
    add_row('Number of Boats Rescued', case.number_of_boats_rescued)
    add_row('Rescuer Names', case.rescuer_name if case.rescuer_name else '')
    add_row('Gist of Rescue', case.gist_of_rescue)

    doc.add_paragraph()

    # Add Rescue Image if available
    if case.rescue_sea_image:
        doc.add_paragraph("Rescue Image:")

        try:
            image_path = case.rescue_sea_image.path
            doc.add_picture(image_path, width=Inches(2.0))  # Adjust width as needed
        except Exception as e:
            doc.add_paragraph(f"Image could not be loaded: {e}")

    # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Rescue_at_Sea_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

#export individual seizure case
@login_required
def seizure_download_view(request, pk):
    case = get_object_or_404(Seizure, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('Seizure Details', level=1)

    # Create table with two columns
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    # Add data
    add_row('Date of Seizure', case.date_of_seizure.strftime('%d-%m-%Y %H:%M') if case.date_of_seizure else '')
    add_row('Place of Seizure', case.place_of_seizure)
    add_row('Seized Item', str(case.seized_item) if case.seized_item else '')
    add_row('Quantity', case.quantity)
    add_row('Lattitude', case.latitude if case.latitude else '')
    add_row('Longitude', case.longitude if case.longitude else '')
    add_row('PS Limit', str(case.ps_limit) if case.ps_limit else '')
    add_row('MPS Limit', str(case.mps_limit) if case.mps_limit else '')
    add_row('Accused', case.accused if case.accused else '')
    add_row('Handed Over To', str(case.handed_over_to) if case.handed_over_to else '')
    add_row('Gist of Seizure', case.gist_of_seizure)

    doc.add_paragraph()

    # Add Seizure Image if available
    if case.seizure_image:
        doc.add_paragraph("Seizure Image:")

        try:
            image_path = case.seizure_image.path
            doc.add_picture(image_path, width=Inches(2.0))  # Adjust width as needed
        except Exception as e:
            doc.add_paragraph(f"Image could not be loaded: {e}")

    # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Seizure_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

#export individual forecast case
@login_required
def forecast_download_view(request, pk):
    case = get_object_or_404(Forecast, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('Forecast Details', level=1)

    # Create table with two columns
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    # Add data
    add_row('Date of Forecast', case.date_of_forecast.strftime('%d-%m-%Y %H:%M') if case.date_of_forecast else '')
    add_row('Place of Forecast', case.place_of_forecast)
    add_row('Forecast Details', case.forecast_details)

    # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Forecast_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

# tn fishermen attack case download view
@login_required
def attack_tnfishermen_download_view(request, pk):
    case = get_object_or_404(AttackOnTNFishermen, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('TN Fishermen Attack Case Details', level=1)
    # Create table with two columns
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''
    # Add data
    add_row('MPS Limit', case.mps_limit)
    add_row('Date of Attack', case.date_of_attack.strftime('%d-%m-%Y %H:%M') if case.date_of_attack else '')
    add_row('Place of Attack', case.place_of_attack)
    add_row('Attacked By', str(case.attacked_by) if case.attacked_by else '')
    add_row('District', case.district)
    add_row('Number of TN Fishermen Injured', str(case.number_of_TNFishermen_injured) if case.number_of_TNFishermen_injured else '')
    add_row('Number of TN Fishermen Missing', str(case.number_of_TNFishermen_missing) if case.number_of_TNFishermen_missing else '')
    add_row('Number of TN Fishermen Died', str(case.number_of_TNFishermen_died) if case.number_of_TNFishermen_died else '')
    add_row('Victims', case.victim_names if case.victim_names else '')
    add_row('Gist of Attack', case.gist_of_attack)
    add_row('trasnfered to', case.transfered_to if case.transfered_to else '')

     # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"TN_Fishermen_Attack_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

# tn fishermen arrest case download view
@login_required
def arrest_tnfishermen_download_view(request, pk):
    case = get_object_or_404(ArrestOfTNFishermen, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('TN Fishermen Arrest Case Details', level=1)
    # Create table with two columns
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''
    # Add data
    add_row('Date of Arrest', case.date_of_arrest.strftime('%d-%m-%Y %H:%M') if case.date_of_arrest else '')
    add_row('Place of Arrest', case.place_of_arrest)
    add_row('Arrested By', str(case.arrested_by) if case.arrested_by else '')
    add_row('District', case.district)
    add_row('MPS Limit', case.mps_limit)
    add_row('Number of TN Fishermen Arrested', str(case.number_of_TNFishermen_arrested) if case.number_of_TNFishermen_arrested else '')
    add_row('fishermen Details', case.arrested_Fishermen_names if case.arrested_Fishermen_names else '')
    add_row('Number of Boat Seized', str(case.no_of_boats_seized) if case.no_of_boats_seized else '')
    add_row('Gist of Arrest', case.gist_of_arrest)
    add_row('Number of TN Fishermen Released', str(case.number_of_TNFishermen_released) if case.number_of_TNFishermen_released else '')
    add_row('Number of Baots Released', case.no_of_boats_released if case.no_of_boats_released else '')

    # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"TN_Fishermen_Arrest_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

# SL fishermen arrest case download view
@login_required
def arrest_slfishermen_download_view(request, pk):
    case = get_object_or_404(ArrestOfSLFishermen, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('SL Fishermen Arrest Case Details', level=1)
    # Create table with two columns
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''
    # Add data
    add_row('Date of Arrest', case.date_of_arrest.strftime('%d-%m-%Y %H:%M') if case.date_of_arrest else '')
    add_row('Place of Arrest', case.place_of_arrest)
    add_row('MPS Limit', case.mps_limit)
    add_row('Arrested By', str(case.arrested_by) if case.arrested_by else '')
    add_row('Number of SL Fishermen Arrested', str(case.number_of_SLFishermen_arrested) if case.number_of_SLFishermen_arrested else '')
    add_row('SL Fishermen Details', case.arrested_Fishermen_name if case.arrested_Fishermen_name else '')
    add_row('Number of Boats Seized', str(case.no_of_boats_seized) if case.no_of_boats_seized else '')
    add_row('Gist of Arrest', case.gist_of_arrest)  

    add_row('Number of SL Fishermen Released', str(case.number_of_SLFishermen_released) if case.number_of_SLFishermen_released else '')
    add_row('Number of Boats Released', case.no_of_boats_released if case.no_of_boats_released else '')

    # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"SL_Fishermen_Arrest_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

# onroad vehicle download view
@login_required
def onroad_vehicle_status_download_view(request, pk):
    case = get_object_or_404(OnRoadVehicleStatus, pk=pk)

    doc = Document()
    doc.add_heading('On Road Vehicle Status', level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    add_row('MPS Limit', str(case.mps_limit))
    add_row('Vehicle Type', case.get_vehicle_type_display())
    add_row('Vehicle Number', case.vehicle_number)
    add_row('Alloted To', case.alloted_to)
    add_row('Working Status', case.get_working_status_display())
    add_row('Remarks', case.remarks if case.remarks else '')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"OnRoad_Vehicle_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

#onwater @login_required
def onwater_vehicle_status_download_view(request, pk):
    case = get_object_or_404(OnWaterVehicleStatus, pk=pk)

    doc = Document()
    doc.add_heading('On Water Vehicle Status', level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    add_row('MPS Limit', str(case.mps_limit))
    add_row('Boat Type', case.get_boat_type_display())
    add_row('Boat Number', case.boat_number if case.boat_number else 'N/A')
    add_row('Working Status', case.get_working_status_display())
    add_row('Remarks', case.remarks if case.remarks else '')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"OnWater_Vehicle_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

#beat download view
@login_required
def beat_download_view(request, pk):
    case = get_object_or_404(BeatDetails, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('Beat Details', level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    add_row('Date of Beat', case.date_of_beat.strftime('%d-%m-%Y') if case.date_of_beat else '')
    add_row('MPS Limit', str(case.mps_limit))
    add_row('Day Beat Count', case.day_beat_count)
    add_row('Night Beat Count', case.night_beat_count)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Beat_Details_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

#proforma download
@login_required
def proforma_download_view(request, pk):
    case = get_object_or_404(Proforma, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('Proforma Summary', level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    add_row('Date of Proforma', case.date_of_proforma.strftime('%d-%m-%Y') if case.date_of_proforma else '')
    add_row('MPS Visited', case.mps_visited)
    add_row('Check Posts Checked', case.check_post_checked)
    add_row('Boat Guard Checked', case.boat_guard_checked)
    add_row('VVC Meetings Conducted', case.vvc_meeting_conducted)
    add_row('Villages Visited', case.villages_visited)
    add_row('Meetings Attended', case.meetings_attended)
    add_row('Awareness Programs Conducted', case.awareness_programs_conducted)
    add_row('Coastal Security Exercises Conducted', case.coastal_security_exercises_conducted)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Proforma_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

#boat patrol download view
@login_required
def boat_patrol_download_view(request, pk):
    case = get_object_or_404(BoatPatrol, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('Boat Patrol Details', level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    # Add data
    add_row('Patrol Officer', case.patrol_officer)
    add_row('Boat Type', case.get_boat_type_display())
    add_row('Boat Number', case.boat_number if case.boat_number else '-')
    add_row('Date of Patrol', case.date_of_patrol.strftime('%d-%m-%Y') if case.date_of_patrol else '')

    start_time = case.patrol_start_time.strftime('%H:%M')
    end_time = case.patrol_end_time.strftime('%H:%M')
    add_row('Time of Patrol', f"{start_time} - {end_time} hrs")

    add_row('Place of Patrol', case.patrol_place)
    add_row('MPS Limit', str(case.mps_limit))
    add_row('No. of Boats Checked', str(case.numberof_boats_checked) if case.numberof_boats_checked else '')
    add_row('Reg. No of Boats Checked', case.registration_numberofboats_checked if case.registration_numberofboats_checked else '-')
    add_row('Remarks', case.remarks if case.remarks else '')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Boat_Patrol_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

# vvc meeting
@login_required
def vvc_download_view(request, pk):
    case = get_object_or_404(VVCmeeting, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('VVC Meeting Details', level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    add_row('Date of VVC', case.date_of_vvc.strftime('%d-%m-%Y') if case.date_of_vvc else '')
    add_row('MPS Limit', str(case.mps_limit))
    add_row('Village Name', case.village_name)
    add_row('No. of Villagers Attended', case.number_of_villagers)
    add_row('Conducted By', case.conducted_by)
    add_row('Details', case.vvc_details if case.vvc_details else '')

    doc.add_paragraph()

    # Add Seizure Image if available
    if case.vvc_image:
        doc.add_paragraph("VVC Image:")

        try:
            image_path = case.vvc_image.path
            doc.add_picture(image_path, width=Inches(2.0))  # Adjust width as needed
        except Exception as e:
            doc.add_paragraph(f"Image could not be loaded: {e}")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"VVC_Meeting_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

# atv doanload view 
@login_required
def atv_patrol_download_view(request, pk):
    case = get_object_or_404(Atvpatrol, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('ATV Patrol Details', level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    # Add data
    add_row('Patrol Officer', case.patrol_officer)
    add_row('ATV Number', case.atv_number if case.atv_number else '-')
    add_row('Date of Patrol', case.date_of_patrol.strftime('%d-%m-%Y') if case.date_of_patrol else '')

    start_time = case.patrol_start_time.strftime('%H:%M')
    end_time = case.patrol_end_time.strftime('%H:%M')
    add_row('Time of Patrol', f"{start_time} - {end_time} hrs")

    add_row('Place of Patrol', case.patrol_place)
    add_row('MPS Limit', str(case.mps_limit))
    add_row('Remarks', case.remarks if case.remarks else '')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"ATV_Patrol_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response


# Vehicle Check Checkpost download view
@login_required
def vehicle_checkpost_download_view(request, pk):
    case = get_object_or_404(VehicleCheckPost, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('Vehicle Check at Check Post Details', level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    # Add data
    add_row('Date of Check', case.date_of_check.strftime('%d-%m-%Y') if case.date_of_check else '')
    
    start_time = case.vehicle_check_start_time.strftime('%H:%M')
    end_time = case.vehicle_check_end_time.strftime('%H:%M')
    add_row('Time of Check', f"{start_time} - {end_time} hrs")

    add_row('Place of Check', str(case.check_post))
    add_row('MPS Limit', str(case.mps_limit))
    add_row('Number of vehicles checked', str(case.number_of_vehicles_checked) if case.number_of_vehicles_checked else '')
    add_row('Reg. No of vehicles checked', str(case.registration_numbers) if case.registration_numbers else '')
    add_row('Remarks', case.remarks if case.remarks else '')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Vehicle_Check_Others_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response

# Vehicle Check Others download view
@login_required
def vehicle_check_others_download_view(request, pk):
    case = get_object_or_404(VehicleCheckothers, pk=pk, user=request.user)

    doc = Document()
    doc.add_heading('Vehicle Check (Others) Details', level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value else ''

    # Add data
    add_row('Date of Check', case.date_of_check.strftime('%d-%m-%Y') if case.date_of_check else '')

    start_time = case.vehicle_check_start_time.strftime('%H:%M')
    end_time = case.vehicle_check_end_time.strftime('%H:%M')
    add_row('Time of Check', f"{start_time} - {end_time} hrs")

    add_row('Place of Check', case.place_of_check)
    add_row('MPS Limit', str(case.mps_limit))
    add_row('Number of vehicles checked', str(case.number_of_vehicles_checked) if case.number_of_vehicles_checked else '')
    add_row('Reg. No of vehicles checked', str(case.registration_numbers) if case.registration_numbers else '')
    add_row('Remarks', case.remarks if case.remarks else '')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Vehicle_Check_Others_{case.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    doc.save(response)
    return response
   
# CSR edit view
@login_required
def csr_edit_view(request, pk):
    csr = get_object_or_404(CSR, pk=pk)
    if request.method == 'POST':
        form = CSRForm(request.POST, instance=csr)
        if form.is_valid():
            form.save()
            messages.success(request, 'CSR updated successfully!')
            return redirect('cases_summary')
    else:
        form = CSRForm(instance=csr)

    return render(request, 'dsr/user/forms/csr_form.html', {'form': form, 'edit_mode': True}) 

# BNSS194 edit view
@login_required
def bnss_missing_edit_view(request, pk):
    case = get_object_or_404(BNSSMissingCase, pk=pk, user=request.user)
    if request.method == 'POST':
        form = BNSSMissingCaseForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            return redirect('cases_summary')
    else:
        form = BNSSMissingCaseForm(instance=case)
    return render(request, 'dsr/user/forms/194bnss_missing_form.html', {'form': form, 'edit_mode': True})

# Other cases edit view
@login_required
def othercases_edit_view(request, pk):
    case = get_object_or_404(OtherCases, pk=pk, user=request.user)
    if request.method == 'POST':
        form = othercasesForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Other case updated successfully!')
            return redirect('cases_summary')
    else:
        form = othercasesForm(instance=case)

    return render(request, 'dsr/user/forms/othercases_form.html', {'form': form, 'edit_mode': True})

# Maritime Act edit view
@login_required
def maritimeact_edit_view(request, pk):
    case = get_object_or_404(MaritimeAct, pk=pk, user=request.user)
    if request.method == 'POST':
        form = MaritimeActForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Maritime Act case updated successfully!')
            return redirect('cases_summary')
    else:
        form = MaritimeActForm(instance=case)

    return render(request, 'dsr/user/forms/maritimeact_form.html', {'form': form, 'edit_mode': True})

# Rescue at Beach edit view
@login_required
def rescue_at_beach_edit_view(request, pk):
    rescue_case = get_object_or_404(RescueAtBeach, pk=pk, user=request.user)
    if request.method == 'POST':
        form = RescueAtBeachForm(request.POST, request.FILES, instance=rescue_case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Rescue at Beach entry updated successfully!')
            return redirect('rescue_seizure_summary')
    else:
        form = RescueAtBeachForm(instance=rescue_case)

    return render(request, 'dsr/user/forms/rescuebeach_form.html', {'form': form, 'edit_mode': True})

# Rescue at Sea edit view
@login_required
def rescue_at_sea_edit_view(request, pk):
    rescue_case = get_object_or_404(RescueAtSea, pk=pk, user=request.user)
    if request.method == 'POST':
        form = RescueAtSeaForm(request.POST, request.FILES, instance=rescue_case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Rescue at Sea entry updated successfully!')
            return redirect('rescue_seizure_summary')
    else:
        form = RescueAtSeaForm(instance=rescue_case)

    return render(request, 'dsr/user/forms/rescueatsea_form.html', {'form': form, 'edit_mode': True})

# Seizure edit view
@login_required
def seizure_edit_view(request, pk):
    seizure_case = get_object_or_404(Seizure, pk=pk, user=request.user)
    if request.method == 'POST':
        form = SeizureForm(request.POST, request.FILES, instance=seizure_case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Seizure entry updated successfully!')
            return redirect('rescue_seizure_summary')
    else:
        form = SeizureForm(instance=seizure_case)

    return render(request, 'dsr/user/forms/seizure_form.html', {'form': form, 'edit_mode': True})

#forecast edit view
@login_required
def forecast_edit_view(request, pk):
    forecast_case = get_object_or_404(Forecast, pk=pk, user=request.user)
    if request.method == 'POST':
        form = ForecastForm(request.POST, instance=forecast_case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Forecast entry updated successfully!')
            return redirect('forecast_summary')
    else:
        form = ForecastForm(instance=forecast_case)

    return render(request, 'dsr/user/forms/forecast_form.html', {'form': form, 'edit_mode': True})

#tn fishermen attack edit view
@login_required
def attack_tnfishermen_edit_view(request, pk):
    case = get_object_or_404(AttackOnTNFishermen, pk=pk, user=request.user)
    if request.method == 'POST':
        form = AttackOnTNFishermenForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'TN Fishermen attack record updated successfully.')
            return redirect('fishermen_attack_arrest_summary')
    else:
        form = AttackOnTNFishermenForm(instance=case)

    return render(request, 'dsr/user/forms/attack_tnfishermen_form.html', {'form': form, 'edit_mode': True})

#tn fishermen arrest edit view
@login_required
def arrest_tnfishermen_edit_view(request, pk):
    case = get_object_or_404(ArrestOfTNFishermen, pk=pk, user=request.user)
    if request.method == 'POST':
        form = ArrestOfTNFishermenForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'TN Fishermen arrest record updated successfully.')
            return redirect('fishermen_attack_arrest_summary')
    else:
        form = ArrestOfTNFishermenForm(instance=case)

    return render(request, 'dsr/user/forms/arrest_tnfishermen_form.html', {'form': form, 'edit_mode': True})

#sl fishermen arrest edit view
@login_required
def arrest_slfishermen_edit_view(request, pk):
    case = get_object_or_404(ArrestOfSLFishermen, pk=pk, user=request.user)
    if request.method == 'POST':
        form = ArrestOfSLFishermenForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'SL Fishermen arrest record updated successfully.')
            return redirect('fishermen_attack_arrest_summary')
    else:
        form = ArrestOfSLFishermenForm(instance=case)

    return render(request, 'dsr/user/forms/arrest_slfishermen_form.html', {'form': form, 'edit_mode': True})

#beat edit
@login_required
def beat_edit_view(request, pk):
    case = get_object_or_404(BeatDetails, pk=pk, user=request.user)
    if request.method == 'POST':
        form = BeatDetailsForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Beat record updated successfully.')
            return redirect('beat__vvc_summary')
    else:
        form = BeatDetailsForm(instance=case)

    return render(request, 'dsr/user/forms/beat_Details_form.html', {'form': form, 'edit_mode': True})

#vvc edit
@login_required
def vvc_edit_view(request, pk):
    case = get_object_or_404(VVCmeeting, pk=pk, user=request.user)
    if request.method == 'POST':
        form = VVCmeetingForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'VVC record updated successfully.')
            return redirect('beat__vvc_summary')
    else:
        form = VVCmeetingForm(instance=case)

    return render(request, 'dsr/user/forms/vvc_meeting_form.html', {'form': form, 'edit_mode': True})

#proforma edit
@login_required
def proforma_edit_view(request, pk):
    case = get_object_or_404(Proforma, pk=pk, user=request.user)
    if request.method == 'POST':
        form = ProformaForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Proforma record updated successfully.')
            return redirect('proforma_summary')
    else:
        form = ProformaForm(instance=case)

    return render(request, 'dsr/user/forms/proforma_form.html', {'form': form, 'edit_mode': True})

#onroad vehicle status edit
@login_required
def onroad_vehicle_status_edit_view(request, pk):
    case = get_object_or_404(OnRoadVehicleStatus, pk=pk)
    if request.method == 'POST':
        form = OnRoadVehicleStatusForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Vehicle Status record updated successfully.')
            return redirect('vehicle_status_summary')
    else:
        form = OnRoadVehicleStatusForm(instance=case)

    return render(request, 'dsr/user/forms/onroad_vehicle_form.html', {'form': form, 'edit_mode': True})

#owater vehicle status edit
@login_required
def onwater_vehicle_status_edit_view(request, pk):
    case = get_object_or_404(OnWaterVehicleStatus, pk=pk)
    if request.method == 'POST':
        form = OnWaterVehicleStatusForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Boat  record updated successfully.')
            return redirect('vehicle_status_summary')
    else:
        form = OnWaterVehicleStatusForm(instance=case)

    return render(request, 'dsr/user/forms/onwater_vehicle_form.html', {'form': form, 'edit_mode': True})

#boat patrol edit
@login_required
def boat_patrol_edit_view(request, pk):
    case = get_object_or_404(BoatPatrol, pk=pk, user=request.user)
    if request.method == 'POST':
        form = BoatPatrolForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Boat Patrol record updated successfully.')
            return redirect('vehicle_check_patrol_summary')
    else:
        form = BoatPatrolForm(instance=case)

    return render(request, 'dsr/user/forms/boat_patrol_form.html', {'form': form, 'edit_mode': True})

#atv patrol  edit
@login_required
def atv_patrol_edit_view(request, pk):
    case = get_object_or_404(Atvpatrol, pk=pk, user=request.user)
    if request.method == 'POST':
        form = AtvpatrolForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'ATV Patrol record updated successfully.')
            return redirect('vehicle_check_patrol_summary')
    else:
        form = AtvpatrolForm(instance=case)

    return render(request, 'dsr/user/forms/atv_patrol_form.html', {'form': form, 'edit_mode': True})

#vehicle check -checkpost edit
@login_required
def vehicle_checkpost_edit_view(request, pk):
    case = get_object_or_404(VehicleCheckPost, pk=pk, user=request.user)
    if request.method == 'POST':
        form = VehicleCheckPostForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Vehicle check  record updated successfully.')
            return redirect('vehicle_check_patrol_summary')
    else:
        form = VehicleCheckPostForm(instance=case)

    return render(request, 'dsr/user/forms/vehicle_checkport_form.html', {'form': form, 'edit_mode': True})

#vehicle check others edit
@login_required
def vehicle_check_others_edit_view(request, pk):
    case = get_object_or_404(VehicleCheckothers, pk=pk, user=request.user)
    if request.method == 'POST':
        form = VehicleCheckothersForm(request.POST, instance=case)
        if form.is_valid():
            form.save()
            messages.success(request, 'Vehicle check record updated successfully.')
            return redirect('vehicle_check_patrol_summary')
    else:
        form = VehicleCheckothersForm(instance=case)

    return render(request, 'dsr/user/forms/vehicle_checkothers_form.html', {'form': form, 'edit_mode': True})

#CSR delete view
@login_required
def csr_delete_view(request, pk):
    csr = get_object_or_404(CSR, pk=pk)
    csr.delete()
    messages.success(request, 'CSR entry deleted successfully!')
    return redirect('cases_summary')  # Adjust this to your summary view name

# BNSS194 delete view
@login_required
def bnss_missing_delete_view(request, pk):
    case = get_object_or_404(BNSSMissingCase, pk=pk, user=request.user)
    case.delete()
    return redirect('cases_summary')

# Other cases delete view
@login_required
def othercases_delete_view(request, pk):
    case = get_object_or_404(OtherCases, pk=pk, user=request.user)
    case.delete()
    messages.success(request, 'Other case entry deleted successfully!')
    return redirect('cases_summary')

# Maritime Act delete view
@login_required
def maritimeact_delete_view(request, pk):
    case = get_object_or_404(MaritimeAct, pk=pk, user=request.user)
    case.delete()
    messages.success(request, 'Maritime Act case entry deleted successfully!')
    return redirect('cases_summary')

# Rescue at Beach delete view
@login_required
def rescue_at_beach_delete_view(request, pk):
    rescue_case = get_object_or_404(RescueAtBeach, pk=pk, user=request.user)
    rescue_case.delete()
    messages.success(request, 'Rescue at Beach entry deleted successfully!')
    return redirect('rescue_seizure_summary')

#rescue at sea delete view
def rescue_at_sea_delete_view(request, pk):
    rescue_case = get_object_or_404(RescueAtSea, pk=pk, user=request.user)
    rescue_case.delete()
    messages.success(request, 'Rescue at Sea entry deleted successfully!')
    return redirect('rescue_seizure_summary')

# Seizure delete view
def seizure_delete_view(request, pk):
    seizure_case = get_object_or_404(Seizure, pk=pk, user=request.user)
    seizure_case.delete()
    messages.success(request, 'Seizure entry deleted successfully!')
    return redirect('rescue_seizure_summary')

#forecast delete view
def forecast_delete_view(request, pk):
    forecast_case = get_object_or_404(Forecast, pk=pk, user=request.user)
    forecast_case.delete()
    messages.success(request, 'Forecast entry deleted successfully!')
    return redirect('forecast_summary')

# fishermen attack_attest delete view
@login_required
def delete_attack_tnfishermen(request, pk):
    attack_case = get_object_or_404(AttackOnTNFishermen, pk=pk, user=request.user)
    attack_case.delete()
    messages.success(request, 'TN Fishermen attack record deleted successfully.')
    return redirect('fishermen_attack_arrest_summary')

@login_required
def delete_arrest_tnfishermen(request, pk):
    obj = get_object_or_404(ArrestOfTNFishermen, pk=pk,user=request.user)
    obj.delete()
    messages.success(request, 'TN Fishermen arrest record deleted successfully.')
    return redirect('fishermen_attack_arrest_summary')

@login_required
def delete_arrest_slfishermen(request, pk):
    obj = get_object_or_404(ArrestOfSLFishermen, pk=pk,user=request.user)
    obj.delete()
    messages.success(request, 'SL Fishermen arrest record deleted successfully.')
    return redirect('fishermen_attack_arrest_summary')

@login_required
def delete_arresbeat_delete_viewt_slfishermen(request, pk):
    obj = get_object_or_404(BeatDetails, pk=pk,user=request.user)
    obj.delete()
    messages.success(request, 'Beat record deleted successfully.')
    return redirect('beat__vvc_summary')

@login_required
def beat_delete_view(request, pk):
    obj = get_object_or_404(BeatDetails, pk=pk,user=request.user)
    obj.delete()
    messages.success(request, 'Beat record deleted successfully.')
    return redirect('beat__vvc_summary')

@login_required
def vvc_delete_view(request, pk):
    obj = get_object_or_404(VVCmeeting, pk=pk,user=request.user)
    obj.delete()
    messages.success(request, 'VVC record deleted successfully.')
    return redirect('beat__vvc_summary')

@login_required
def proforma_delete_view(request, pk):
    obj = get_object_or_404(Proforma, pk=pk,user=request.user)
    obj.delete()
    messages.success(request, 'Proforma record deleted successfully.')
    return redirect('proforma_summary')

@login_required
def onroad_vehicle_status_delete_view(request, pk):
    obj = get_object_or_404(OnRoadVehicleStatus, pk=pk)
    obj.delete()
    messages.success(request, 'On Road Vehicle Status record deleted successfully.')
    return redirect('vehicle_status_summary')

@login_required
def onwater_vehicle_status_delete_view(request, pk):
    obj = get_object_or_404(OnWaterVehicleStatus, pk=pk)
    obj.delete()
    messages.success(request, 'Boat  Status record deleted successfully.')
    return redirect('vehicle_status_summary')

@login_required
def boat_patrol_delete_view(request, pk):
    obj = get_object_or_404(BoatPatrol, pk=pk,user=request.user)
    obj.delete()
    messages.success(request, 'Boat  Patrol record deleted successfully.')
    return redirect('vehicle_check_patrol_summary')

@login_required
def atv_patrol_delete_view(request, pk):
    obj = get_object_or_404(Atvpatrol, pk=pk,user=request.user)
    obj.delete()
    messages.success(request, 'ATV  Patrol record deleted successfully.')
    return redirect('vehicle_check_patrol_summary')

@login_required
def vehicle_checkpost_delete_view(request, pk):
    obj = get_object_or_404(VehicleCheckPost, pk=pk,user=request.user)
    obj.delete()
    messages.success(request, 'Vehicle Check  record deleted successfully.')
    return redirect('vehicle_check_patrol_summary')

@login_required
def vehicle_check_others_delete_view(request, pk):
    obj = get_object_or_404(VehicleCheckothers, pk=pk,user=request.user)
    obj.delete()
    messages.success(request, 'Vehicle Check  record deleted successfully.')
    return redirect('vehicle_check_patrol_summary')


